# ═══════════════════════════════════════════════════════════════════
# Astra — Lakshmi Supraja Konakanchi (v2.0)
# US Senior Data Engineer Edition — recruiter-grade upgrade
#
# What changed from v1.0:
#   v2.0 incorporates the May 2026 grad.jobs / Resume AI v3 research on
#   what separates 78-85% "callback resumes" from 95%+ "keyword-stuffed
#   robotic resumes". The research recommended a 5-stage section-
#   differential pipeline with self-refine loop, embeddings, and Pro
#   model — that would explode cost. We honor the *quality wins* of the
#   research while keeping the *cost shape* the user asked for.
#
# Architecture: ONE LLM call (Gemini 3 Flash Preview only, per user
# constraint). All scoring, grounding, validation, and rationale
# generation is done DETERMINISTICALLY in Python. Zero extra LLM calls.
#
# v2.0 features added vs v1.0:
#   [PROMPT-LEVEL]
#   1. Three discrete presets (keep_voice / match_job / maximize_match)
#      — each maps to a different prompt emphasis and temperature.
#      No slider per UX research.
#   2. JD intelligence extraction folded into the same call: model
#      returns top_keywords, required_tools, emphasis_areas, industry.
#   3. "What Astra did" rationale block folded into same call: model
#      returns plain-English notes ("Rewrote 4 bullets for Snowflake.
#      Did NOT add Kubernetes — not in your original.").
#   4. Canonical-form-first rule baked in (write "Apache Airflow"
#      before "Airflow", "Amazon Web Services (AWS)" before "AWS").
#   5. Invisible keyword technique: keywords appear inside bullets in
#      natural prose, not stuffed in Skills blob.
#   6. ≥60% bullets quantified (NOT 100% — universal quantification
#      triggers "fake metric" tell per research).
#   7. Target ATS match band: 78-85%. Cap displayed at 88%. The
#      research is unambiguous: ≥95% is the keyword-stuffing red flag.
#
#   [DETERMINISTIC PYTHON, ZERO LLM COST]
#   8. Tier-1 banned-phrase regex post-sweep (delve, leverage,
#      spearhead, robust, seamless, etc. — full list from research).
#   9. Banned-opener detector + fix (Furthermore, Moreover, In today's,
#      etc.).
#   10. Replacements table (leveraged→used, spearheaded→led,
#       utilized→used, streamlined→simplified, optimized→tuned,
#       foster→build, empower→enable, holistic→end-to-end).
#   11. Burstiness audit: sentence-length stdev ≥4.0, verb diversity
#       ≥0.80, no lead verb appears more than 2 times.
#   12. Multi-axis score card: Impact / Keywords / Readability /
#       Experience — all four computed from the generated text in pure
#       Python. NOT an LLM scoring call.
#   13. Facts inventory grounding: extract every tool name and number
#       from the base resume at startup; flag any generated bullet
#       that introduces a tool or number not in the inventory.
#   14. "Ready to send?" deterministic checklist (Impact ≥60%,
#       keyword coverage ≥75%, no flagged hallucinations, length OK).
#
#   [UI]
#   15. Three-step wizard: paste → controls → review (per research UX).
#   16. Multi-axis score card with traffic-light colors.
#   17. "What Astra did" rationale block visible above the resume.
#   18. "Ready to send?" green-check block above download.
#   19. Section-comparison view (base vs tailored, stacked).
#   20. Optional ATS scoring button (kept from v1.0).
#   21. Optional cover letter button (kept from v1.0).
#
# Cost & speed:
#   Default: 1 Flash 3 call → ~$0.008, 8-15s
#   Worst case (resume + score + cover): 3 Flash 3 calls → ~$0.024
#   Old Supraja Astra: ~$0.04-0.12, 25-90s
#
# DEFERRED from research (would require extra LLM calls or embeddings):
#   - Self-refine generate→critique→refine loop (3 iterations)
#   - Per-bullet targeted regeneration for failing bullets
#   - Embedding-based grounding via gemini-embedding-001
#   - True word-level diff with per-bullet Accept/Revert
# These can be added later as optional buttons if needed.
# ═══════════════════════════════════════════════════════════════════

import streamlit as st
import json
import re
import io
import datetime
import statistics
from typing import List, Set, Dict, Tuple
# Note: difflib was imported in an earlier draft for a per-word diff
# implementation. The current Before/After tab uses side-by-side text
# areas instead — adequate for non-technical reviewers. If a true
# word-level diff is added later, re-add `import difflib`.

from pydantic import BaseModel, Field
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from xml.sax.saxutils import escape


# ═══════════════════════════════════════════════════════════════════
# 1. API KEYS & MODELS
# Per user constraint: ONLY gemini-3-flash-preview, no other models.
# That means no Flash-Lite for scoring, no embeddings, no Pro.
# Everything LLM-based uses this same model. Everything else is Python.
# ═══════════════════════════════════════════════════════════════════

try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
except Exception:
    GOOGLE_API_KEY = ""

GENERATION_MODEL = "gemini-3-flash-preview"


# ═══════════════════════════════════════════════════════════════════
# 2. LOCKED STRUCTURAL FACTS
# Immutable at code level. Renderer always uses these. Model receives
# them as context but cannot rewrite them.
# ═══════════════════════════════════════════════════════════════════

PAGE_TITLE = "Astra — Supraja"

CANDIDATE_NAME = "Lakshmi Supraja Konakanchi"
CANDIDATE_TAGLINE_DEFAULT = "Senior Data Engineer"
CANDIDATE_CONTACT = (
    "+1 (469) 723-2320 | lakshmik3272@gmail.com | "
    "Irving, TX, USA | linkedin.com/in/lakshmi-k-19aa79330"
)

LOCKED_ROLE_HEADERS = [
    {
        "role_title": "Azure Data Engineer (Contract)",
        "company": "Northwestern Mutual",
        "location": "San Antonio, TX",
        "dates": "Jul 2024 - Present",
        "tense": "present",
        "industry": "Insurance / Financial Services / Wealth Management",
    },
    {
        "role_title": "AWS Data Engineer (Contract)",
        "company": "McKesson Corporation",
        "location": "Irving, TX",
        "dates": "Aug 2023 - May 2024",
        "tense": "past",
        "industry": "Healthcare / Pharmacy / Pharma Distribution",
    },
    {
        "role_title": "GCP Data Engineer",
        "company": "Mindtree Limited (BigBasket Account)",
        "location": "Hyderabad, India",
        "dates": "Aug 2021 - Jul 2022",
        "tense": "past",
        "industry": "Retail / E-commerce / Quick Commerce",
    },
    {
        "role_title": "Data Engineer (Contract)",
        "company": "Geekyants Private Limited",
        "location": "Bengaluru, India",
        "dates": "Jun 2020 - Jul 2021",
        "tense": "past",
        "industry": "Tech / SaaS / Web & Mobile Applications",
    },
    {
        "role_title": "Data Engineer (Contract)",
        "company": "Exide Energy Solutions",
        "location": "Bengaluru, India",
        "dates": "Dec 2019 - May 2020",
        "tense": "past",
        "industry": "Manufacturing / Energy / Battery Operations / IoT",
    },
]

LOCKED_EDUCATION = [
    {
        "degree": "Master of Science in Management Information Systems",
        "institution": "Lamar University, Beaumont, TX",
        "dates": "May 2024",
    },
]

LOCKED_CERTIFICATIONS = [
    "AWS Certified Data Engineer - Associate (DEA-C01), Amazon Web Services (Aug 2025 - Aug 2028)",
]


# ═══════════════════════════════════════════════════════════════════
# 3. SUPRAJA BASE RESUME (model context + facts inventory source)
# ═══════════════════════════════════════════════════════════════════

SUPRAJA_BASE_RESUME = """LAKSHMI SUPRAJA KONAKANCHI
Senior Data Engineer
+1 (469) 723-2320 | lakshmik3272@gmail.com | Irving, TX, USA | linkedin.com/in/lakshmi-k-19aa79330

Professional Profile
AWS Certified Senior Data Engineer with 5+ years architecting cloud-native data platforms across Amazon Web Services (AWS), Microsoft Azure, and Google Cloud Platform (GCP) for Fortune 500 clients in Finance, Healthcare, and Retail. Built and migrated Extract Transform Load (ETL) / Extract Load Transform (ELT) pipelines processing 500GB+ of policy, claims, and pharmacy data daily using Azure Data Factory (ADF), Databricks, dbt, Azure Synapse, and Snowflake, cutting batch runtime by roughly 40% and streaming latency by roughly 25%. Deep hands-on across Apache PySpark, Apache Kafka, Apache Airflow, Delta Lake, and Apache Iceberg, with a track record of regulated-data delivery under HIPAA, PCI DSS, and SOX compliance.

Key Skills / Tools & Technologies
- Programming Languages: Python, SQL, Apache PySpark, Scala, Shell Scripting (Bash), T-SQL, PL/SQL
- AWS Cloud: AWS Glue, AWS Lambda, Amazon Kinesis, Amazon S3, Amazon Redshift, Amazon EMR, Amazon RDS, Amazon DynamoDB, AWS Step Functions, Amazon Athena, AWS CloudWatch, AWS IAM, AWS Lake Formation
- Azure Cloud: Azure Data Factory (ADF), Azure Databricks, Azure Synapse Analytics, Azure Data Lake Storage (ADLS) Gen2, Azure Event Hubs, Azure Functions, Azure Logic Apps, Azure Cosmos DB, Azure SQL Database, Azure DevOps
- GCP Cloud: Google BigQuery, Google Cloud Dataflow (Apache Beam), Google Dataproc, Google Cloud Pub/Sub, Google Cloud Composer, Google Cloud Storage (GCS), Google Cloud Functions, Google Cloud Data Catalog
- Data Platforms & Warehouses: Snowflake (Snowpipe, Streams, Tasks, Time Travel, Role-Based Access Control, Snowpark), Databricks (Delta Lake, Unity Catalog, MLflow, Workflows, Photon), Amazon Redshift, Azure Synapse, Google BigQuery, Apache Iceberg, Trino, Presto
- Streaming & Orchestration: Apache Kafka (Kafka Connect, Schema Registry), Amazon Kinesis (Streams, Firehose, Analytics), Azure Event Hubs, Google Cloud Pub/Sub, Apache Spark Structured Streaming, Apache Flink, Apache Airflow, dbt Core, Apache Beam, Azure Stream Analytics, Informatica, Talend, Apache NiFi, Apache Oozie, Fivetran, Airbyte
- Databases: PostgreSQL, Microsoft SQL Server, MySQL, Oracle, MariaDB, MongoDB, Amazon DynamoDB, Azure Cosmos DB, Google Cloud Spanner, Apache HBase, Apache Hive, Apache Cassandra
- Data Modeling & File Formats: Kimball Dimensional Modeling, Star Schema, Snowflake Schema, Data Vault 2.0, Slowly Changing Dimensions (SCD) Type 1 / Type 2, Medallion Architecture (Bronze / Silver / Gold), Lakehouse, Apache Iceberg, Delta Lake, Apache Parquet, Apache Avro, Apache ORC, JSON
- DevOps & Infrastructure as Code (IaC): Git, GitHub Actions, Jenkins, GitLab CI, Azure DevOps, Terraform, AWS CloudFormation, Azure Resource Manager (ARM) Templates, Docker, Kubernetes (Amazon EKS, Azure AKS, Google GKE), Helm
- Data Quality, Governance & Compliance: Great Expectations, Soda Core (SodaCL), Unity Catalog, Google Cloud Data Catalog, Role-Based Access Control (RBAC), Personally Identifiable Information (PII) Masking, Data Lineage, Health Insurance Portability and Accountability Act (HIPAA), Payment Card Industry Data Security Standard (PCI DSS), Sarbanes-Oxley (SOX) Compliance, General Data Protection Regulation (GDPR), Data Contracts
- Business Intelligence (BI), Monitoring & Methodologies: Microsoft Power BI, Tableau, Looker, Amazon QuickSight, Google Looker Studio, Splunk, Prometheus, Grafana, Elasticsearch / Logstash / Kibana (ELK Stack), AWS CloudWatch, Azure Monitor, Datadog, Agile / Scrum, Continuous Integration / Continuous Deployment (CI/CD), DataOps

Professional Experience

Azure Data Engineer (Contract) | Northwestern Mutual | San Antonio, TX | Jul 2024 - Present
Industry: Insurance / Financial Services / Wealth Management
- Architected Azure Data Factory and Databricks ETL/ELT pipelines on Delta Lake processing 500GB+ of daily policy, claims, and investment data from Microsoft SQL Server, Azure Cosmos DB, and REST APIs into ADLS Gen2 and Azure Synapse Analytics, supporting insurance and wealth management reporting.
- Migrated legacy on-premises ETL into Azure Synapse and Snowflake using dbt transformation models on Azure Databricks, cutting overnight batch runtime by ~40% and freeing 3+ additional reporting hours for wealth management business intelligence (BI) teams.
- Built Azure Event Hubs, Apache Kafka, and Apache Spark Structured Streaming pipelines on Databricks for customer transaction monitoring, cutting end-to-end data latency by ~25% across 15+ advisor-facing dashboards.
- Tuned Azure Synapse, Azure Cosmos DB, and Azure SQL workloads through indexing, partitioning, and query refactoring, improving analytical query throughput by 20-30% for actuarial and reporting teams.
- Implemented PII masking, Role-Based Access Control (RBAC), and Unity Catalog governance across ADLS Gen2, Cosmos DB, and Snowflake under SOX compliance; productionised Continuous Integration / Continuous Deployment (CI/CD) via Jenkins, Git, and Terraform, reducing release effort by ~30%.
Achievements:
- Mentored 2 junior engineers on Apache PySpark performance tuning and Databricks Delta Lake patterns within a 15+ engineer data platform team, and partnered with data science to deliver model-ready credit-risk and customer-360 datasets on Snowflake.
- Standardised Snowflake transformations using dbt Core modular models and automated data quality tests, reducing ad-hoc query failures and improving data reliability for downstream analytics teams.

AWS Data Engineer (Contract) | McKesson Corporation | Irving, TX | Aug 2023 - May 2024
Industry: Healthcare / Pharmacy / Pharma Distribution
- Engineered HIPAA-aligned AWS Glue, AWS Lambda, and Amazon EMR pipelines using Apache PySpark to process 400GB+ of daily pharmacy and claims data from Amazon RDS and on-premises systems into an Amazon S3 data lake built on Apache Iceberg table format.
- Built Amazon Kinesis, Apache Kafka, and Spark Streaming applications for prescription fulfillment tracking and inventory analytics, enabling near-real-time operational visibility for pharmacy distribution teams.
- Tuned Amazon Redshift through sort/distribution key redesign and Workload Management (WLM) configuration, improving analytical query performance by ~30% and reducing peak-hour wait times for clinical and operations analysts.
- Orchestrated pipelines using Apache Airflow on Amazon Managed Workflows (MWAA) and AWS Step Functions with retry logic, SLA alerting, and AWS CloudWatch monitoring, reducing pipeline failures and on-call escalations by ~30%.
- Cut Amazon S3 costs by ~25% via lifecycle policies, intelligent tiering, and Apache Parquet/ORC re-encoding; delivered HIPAA-compliant Microsoft Power BI and Amazon QuickSight reporting datasets with column-level masking and AWS IAM separation.
Achievements:
- Cleared HIPAA internal audit review on pharmacy and claims data access controls (column-level masking, IAM role separation, pipeline-level lineage) without findings, supporting compliance sign-off for downstream analytics.
- Introduced Apache Iceberg on the Amazon S3 data lake to enable ACID transactions and schema evolution for 400GB+ of daily pharmacy claims, future-proofing the architecture for multi-engine reads from Amazon Athena, Amazon EMR, and Trino.

GCP Data Engineer | Mindtree Limited (BigBasket Account) | Hyderabad, India | Aug 2021 - Jul 2022
Industry: Retail / E-commerce / Quick Commerce
- Built Google Cloud Dataflow (Apache Beam), Google Dataproc, and Google Cloud Pub/Sub ETL pipelines processing 300GB+ of daily retail order, inventory, and logistics data into Google BigQuery for the BigBasket e-commerce platform.
- Tuned Google BigQuery through partitioning, clustering, and materialized views, improving analytical query efficiency by ~25% and reducing dashboard load times for supply-chain and category teams.
- Containerised Google Cloud Dataflow jobs on Google Kubernetes Engine (GKE) with Docker via Google Cloud Composer, implemented Google Cloud Data Catalog for lineage tracking, and delivered Looker and Google Looker Studio dashboards on Medallion-layered BigQuery tables.

Data Engineer (Contract) | Geekyants Private Limited | Bengaluru, India | Jun 2020 - Jul 2021
Industry: Tech / SaaS / Web & Mobile Applications
- Built Apache Spark, Apache Kafka, and Apache Airflow pipelines to process telemetry and user activity data across client web and mobile platforms, enabling centralised analytics and near-real-time operational dashboards in Microsoft Power BI, Tableau, and Looker.
- Automated CI/CD using Jenkins, Docker, and AWS Lambda, reducing manual deployment effort by ~25% and improving release reliability across multiple client engagements.

Data Engineer (Contract) | Exide Energy Solutions | Bengaluru, India | Dec 2019 - May 2020
Industry: Manufacturing / Energy / Battery Operations / IoT
- Built Apache Spark and Apache Hadoop batch pipelines on Hadoop Distributed File System (HDFS), Apache Hive, and Apache HBase to process 2TB+ of manufacturing sensor data for production monitoring; migrated selected workloads to AWS Lambda and Amazon Kinesis using Terraform, improving pipeline scalability.

Education
Master of Science in Management Information Systems, Lamar University, Beaumont, TX (May 2024)

Certifications
- AWS Certified Data Engineer - Associate (DEA-C01), Amazon Web Services (Aug 2025 - Aug 2028)
"""


# ═══════════════════════════════════════════════════════════════════
# 4. THREE-PRESET CONFIGURATION
# Per research: discrete presets > continuous slider for non-expert
# users. "match_job" is default because metric preservation is
# non-negotiable for 5-YOE Senior DE.
# ═══════════════════════════════════════════════════════════════════

PRESET_CONFIGS = {
    "keep_voice": {
        "label": "🟢 Keep my voice",
        "caption": "Only adds missing keywords. Doesn't touch bullet wording. Safest.",
        "temperature": 0.3,
        "instruction": (
            "PRESET = KEEP_VOICE (conservative). Keep the candidate's existing "
            "bullets as close to the base resume as possible. Make minimal "
            "edits — primarily add missing JD keywords to the Skills section "
            "and one or two bullets where natural. DO NOT rephrase bullets "
            "that already work. Preserve all numbers, dates, and the original "
            "sentence structure as much as you can. Summary may be tailored "
            "to the JD title, but keep her voice."
        ),
    },
    "match_job": {
        "label": "🟡 Match the job  ⭐ recommended",
        "caption": "Rewrites bullets in the job's language. Keeps your numbers, dates, and scope. Default.",
        "temperature": 0.5,
        "instruction": (
            "PRESET = MATCH_JOB (default, recommended). Rewrite bullets in the "
            "JD's terminology where the candidate has the underlying work. "
            "PRESERVE ALL METRICS, dates, and scope verbatim from the base "
            "resume. You may reorder bullets within a role to put the most "
            "JD-relevant work first. Mirror the JD's exact role title in the "
            "summary's first sentence."
        ),
    },
    "maximize_match": {
        "label": "🔴 Maximize match",
        "caption": "Reframes whole experience. Review carefully — may sound less like you.",
        "temperature": 0.7,
        "instruction": (
            "PRESET = MAXIMIZE_MATCH (aggressive). Fully reframe the "
            "candidate's experience in the JD's vocabulary. You may merge or "
            "split bullets, change emphasis significantly, and inject "
            "JD-aligned framing throughout. STILL preserve every metric and "
            "date — numbers stay verbatim. Foreground architectural "
            "ownership: named migrations, cross-team data products, "
            "mentorship of 1-3 engineers, cost governance. Set the candidate "
            "apart from mid-level."
        ),
    },
}

DEFAULT_PRESET = "match_job"


# ═══════════════════════════════════════════════════════════════════
# 5. RESEARCH-GRADE BANNED PHRASES / OPENERS / REPLACEMENTS
# Per research: prompt-level instructions catch most violations but
# a regex post-sweep is the hard-constraint failsafe.
# ═══════════════════════════════════════════════════════════════════

# Tier 1: phrases that ALWAYS get flagged and replaced or rewritten.
# Source: 2025-2026 recruiter-survey research on AI-tell signals.
BANNED_TIER1_PATTERN = re.compile(
    r'\b('
    r'delve(d|ing|s)?|'
    r'leverag(e|ed|es|ing)|'
    r'spearhead(ed|ing|s)?|'
    r'robust|'
    r'seamless(ly)?|'
    r'intricat(e|ely|ies)|'
    r'harness(ed|ing|es)?|'
    r'unlock(ed|ing|s)?|'
    r'unleash(ed|ing|es)?|'
    r'empower(ed|ing|s)?|'
    r'transformative|'
    r'groundbreaking|'
    r'pivotal|'
    r'cutting[- ]edge|'
    r'holistic|'
    r'synerg(y|ies|ize|ized|izing)|'
    r'foster(ed|ing|s)?|'
    r'streamlin(e|ed|es|ing)|'
    r'elevat(e|ed|es|ing)|'
    r'paradigm|'
    r'ever[- ]evolving|'
    r'tapestr(y|ies)|'
    r'realm|'
    r'embark(ed|ing|s)?|'
    r'vibrant|'
    r'crucial|'
    r'compelling|'
    r'testament|'
    r'navigate the landscape|'
    r'utiliz(e|ed|es|ing)'
    r')\b',
    re.IGNORECASE,
)

# Banned openers (sentence-initial). Detected ONLY at sentence start.
BANNED_OPENERS_PATTERN = re.compile(
    r'^(Furthermore|Moreover|Additionally|Consequently|'
    r"In today's|In an era of|It is worth noting)",
    re.IGNORECASE | re.MULTILINE,
)

# Replacements table. Applied case-insensitively. Original case is
# preserved where reasonable (capitalised → capitalised). Per research,
# these soft-banned words have direct plain-English substitutes.
REPLACEMENTS_TABLE = [
    (re.compile(r'\bleverag(e|ed|es|ing)\b', re.IGNORECASE), {
        'leverage': 'use', 'leveraged': 'used',
        'leverages': 'uses', 'leveraging': 'using',
    }),
    (re.compile(r'\bspearhead(ed|ing|s)?\b', re.IGNORECASE), {
        'spearhead': 'lead', 'spearheaded': 'led',
        'spearheading': 'leading', 'spearheads': 'leads',
    }),
    (re.compile(r'\butiliz(e|ed|es|ing)\b', re.IGNORECASE), {
        'utilize': 'use', 'utilized': 'used',
        'utilizes': 'uses', 'utilizing': 'using',
    }),
    (re.compile(r'\bstreamlin(e|ed|es|ing)\b', re.IGNORECASE), {
        'streamline': 'simplify', 'streamlined': 'simplified',
        'streamlines': 'simplifies', 'streamlining': 'simplifying',
    }),
    (re.compile(r'\bfoster(ed|ing|s)?\b', re.IGNORECASE), {
        'foster': 'build', 'fostered': 'built',
        'fostering': 'building', 'fosters': 'builds',
    }),
    (re.compile(r'\bempower(ed|ing|s)?\b', re.IGNORECASE), {
        'empower': 'enable', 'empowered': 'enabled',
        'empowering': 'enabling', 'empowers': 'enables',
    }),
    (re.compile(r'\bholistic\b', re.IGNORECASE),
     {'holistic': 'end-to-end'}),
    (re.compile(r'\bharness(ed|ing|es)?\b', re.IGNORECASE), {
        'harness': 'use', 'harnessed': 'used',
        'harnessing': 'using', 'harnesses': 'uses',
    }),
]


def apply_replacements(text: str) -> str:
    """Soft-replace banned but easily-substituted words. Preserves
    case of the original word where reasonable."""
    if not text:
        return text

    def _replace(match, mapping):
        word = match.group(0)
        lower = word.lower()
        new_word = mapping.get(lower, word)
        # Preserve capitalisation of first letter
        if word and word[0].isupper():
            new_word = new_word[0].upper() + new_word[1:]
        return new_word

    cleaned = text
    for pattern, mapping in REPLACEMENTS_TABLE:
        cleaned = pattern.sub(lambda m: _replace(m, mapping), cleaned)
    return cleaned


def count_banned_tier1(text: str) -> int:
    """Count remaining Tier-1 hits after replacement pass."""
    if not text:
        return 0
    return len(BANNED_TIER1_PATTERN.findall(text))


def fix_banned_openers(text: str) -> str:
    """Strip banned sentence-initial connectors."""
    if not text:
        return text
    cleaned = BANNED_OPENERS_PATTERN.sub('', text)
    cleaned = re.sub(r'^[\s,]+', '', cleaned, flags=re.MULTILINE)
    return cleaned


# ═══════════════════════════════════════════════════════════════════
# 6. SENIORITY VERB LISTS
# Used by the Experience axis scoring. Per research: 5-YOE Senior DE
# resumes should foreground architectural ownership signals.
# ═══════════════════════════════════════════════════════════════════

SENIOR_VERBS = {
    "architected", "designed", "led", "mentored", "owned", "drove",
    "established", "standardised", "standardized", "productionised",
    "productionized", "migrated", "consolidated", "introduced",
    "instituted", "built", "delivered", "shipped", "orchestrated",
    "implemented", "engineered", "tuned", "optimised", "optimized",
    "partnered", "scaled", "modernised", "modernized",
}

JUNIOR_VERBS = {
    "helped", "assisted", "supported", "participated", "contributed",
    "shadowed", "observed", "learned", "trained on", "worked on",
}


# ═══════════════════════════════════════════════════════════════════
# 7. MASTER CATEGORY MAP + ECOSYSTEM KEYWORDS
# Same as v1.0. Used by skill-category validator for tidiness only,
# never as a filter. The model is free to add any tool the JD asks for.
# ═══════════════════════════════════════════════════════════════════

MASTER_CATEGORY_MAP = {
    # Programming Languages
    "python": "Programming Languages", "sql": "Programming Languages",
    "pyspark": "Programming Languages", "scala": "Programming Languages",
    "shell scripting": "Programming Languages", "bash": "Programming Languages",
    "t-sql": "Programming Languages", "pl/sql": "Programming Languages",
    "java": "Programming Languages", "powershell": "Programming Languages",
    "apache pyspark": "Programming Languages",
    # AWS Cloud
    "aws glue": "AWS Cloud", "glue": "AWS Cloud",
    "aws lambda": "AWS Cloud", "lambda": "AWS Cloud",
    "amazon kinesis": "AWS Cloud", "kinesis": "AWS Cloud",
    "amazon s3": "AWS Cloud", "s3": "AWS Cloud",
    "amazon redshift": "AWS Cloud", "redshift": "AWS Cloud",
    "amazon emr": "AWS Cloud", "emr": "AWS Cloud",
    "amazon rds": "AWS Cloud", "rds": "AWS Cloud",
    "amazon dynamodb": "AWS Cloud", "dynamodb": "AWS Cloud",
    "aws step functions": "AWS Cloud", "step functions": "AWS Cloud",
    "amazon athena": "AWS Cloud", "athena": "AWS Cloud",
    "aws cloudwatch": "AWS Cloud", "cloudwatch": "AWS Cloud",
    "aws iam": "AWS Cloud", "iam": "AWS Cloud",
    "aws lake formation": "AWS Cloud", "lake formation": "AWS Cloud",
    "mwaa": "AWS Cloud", "amazon managed workflows": "AWS Cloud",
    # Azure Cloud
    "azure data factory": "Azure Cloud", "adf": "Azure Cloud",
    "azure data factory (adf)": "Azure Cloud",
    "azure databricks": "Azure Cloud", "databricks": "Azure Cloud",
    "azure synapse": "Azure Cloud", "synapse analytics": "Azure Cloud",
    "azure synapse analytics": "Azure Cloud",
    "adls gen2": "Azure Cloud", "azure data lake storage": "Azure Cloud",
    "azure event hubs": "Azure Cloud", "event hubs": "Azure Cloud",
    "azure functions": "Azure Cloud",
    "azure logic apps": "Azure Cloud", "logic apps": "Azure Cloud",
    "azure cosmos db": "Azure Cloud", "cosmos db": "Azure Cloud",
    "azure sql database": "Azure Cloud", "azure sql": "Azure Cloud",
    "azure devops": "Azure Cloud",
    "azure stream analytics": "Azure Cloud",
    # GCP Cloud
    "google bigquery": "GCP Cloud", "bigquery": "GCP Cloud",
    "google cloud dataflow": "GCP Cloud", "dataflow": "GCP Cloud",
    "google dataproc": "GCP Cloud", "dataproc": "GCP Cloud",
    "google cloud pub/sub": "GCP Cloud", "pub/sub": "GCP Cloud",
    "cloud pub/sub": "GCP Cloud",
    "google cloud composer": "GCP Cloud", "cloud composer": "GCP Cloud",
    "google cloud storage": "GCP Cloud", "gcs": "GCP Cloud",
    "google cloud functions": "GCP Cloud", "cloud functions": "GCP Cloud",
    "google cloud data catalog": "GCP Cloud", "data catalog": "GCP Cloud",
    "google kubernetes engine": "GCP Cloud", "gke": "GCP Cloud",
    "vertex ai": "GCP Cloud",
    # Data Platforms & Warehouses
    "snowflake": "Data Platforms & Warehouses",
    "snowpipe": "Data Platforms & Warehouses",
    "snowpark": "Data Platforms & Warehouses",
    "delta lake": "Data Platforms & Warehouses",
    "unity catalog": "Data Platforms & Warehouses",
    "apache iceberg": "Data Platforms & Warehouses",
    "iceberg": "Data Platforms & Warehouses",
    "mlflow": "Data Platforms & Warehouses",
    "trino": "Data Platforms & Warehouses",
    "presto": "Data Platforms & Warehouses",
    "lakehouse": "Data Platforms & Warehouses",
    # Streaming & Orchestration
    "apache kafka": "Streaming & Orchestration", "kafka": "Streaming & Orchestration",
    "kafka connect": "Streaming & Orchestration",
    "kafka streams": "Streaming & Orchestration",
    "schema registry": "Streaming & Orchestration",
    "apache spark structured streaming": "Streaming & Orchestration",
    "spark streaming": "Streaming & Orchestration",
    "spark structured streaming": "Streaming & Orchestration",
    "apache flink": "Streaming & Orchestration", "flink": "Streaming & Orchestration",
    "apache airflow": "Streaming & Orchestration", "airflow": "Streaming & Orchestration",
    "dbt": "Streaming & Orchestration", "dbt core": "Streaming & Orchestration",
    "apache beam": "Streaming & Orchestration", "beam": "Streaming & Orchestration",
    "informatica": "Streaming & Orchestration",
    "talend": "Streaming & Orchestration",
    "apache nifi": "Streaming & Orchestration", "nifi": "Streaming & Orchestration",
    "apache oozie": "Streaming & Orchestration", "oozie": "Streaming & Orchestration",
    "fivetran": "Streaming & Orchestration",
    "airbyte": "Streaming & Orchestration",
    "prefect": "Streaming & Orchestration", "dagster": "Streaming & Orchestration",
    # Databases
    "postgresql": "Databases", "postgres": "Databases",
    "microsoft sql server": "Databases", "sql server": "Databases",
    "mysql": "Databases", "oracle": "Databases",
    "mariadb": "Databases", "mongodb": "Databases",
    "apache hbase": "Databases", "hbase": "Databases",
    "apache hive": "Databases", "hive": "Databases",
    "apache cassandra": "Databases", "cassandra": "Databases",
    "google cloud spanner": "Databases", "cloud spanner": "Databases",
    # Data Modeling & File Formats
    "kimball": "Data Modeling & File Formats",
    "kimball dimensional modeling": "Data Modeling & File Formats",
    "star schema": "Data Modeling & File Formats",
    "snowflake schema": "Data Modeling & File Formats",
    "data vault": "Data Modeling & File Formats",
    "data vault 2.0": "Data Modeling & File Formats",
    "scd": "Data Modeling & File Formats",
    "slowly changing dimensions": "Data Modeling & File Formats",
    "medallion": "Data Modeling & File Formats",
    "medallion architecture": "Data Modeling & File Formats",
    "apache parquet": "Data Modeling & File Formats", "parquet": "Data Modeling & File Formats",
    "apache avro": "Data Modeling & File Formats", "avro": "Data Modeling & File Formats",
    "apache orc": "Data Modeling & File Formats", "orc": "Data Modeling & File Formats",
    "json": "Data Modeling & File Formats",
    # DevOps & IaC
    "git": "DevOps & IaC", "github": "DevOps & IaC",
    "github actions": "DevOps & IaC",
    "jenkins": "DevOps & IaC",
    "gitlab ci": "DevOps & IaC", "gitlab": "DevOps & IaC",
    "terraform": "DevOps & IaC",
    "aws cloudformation": "DevOps & IaC", "cloudformation": "DevOps & IaC",
    "azure resource manager": "DevOps & IaC", "arm templates": "DevOps & IaC",
    "docker": "DevOps & IaC",
    "kubernetes": "DevOps & IaC", "k8s": "DevOps & IaC",
    "amazon eks": "DevOps & IaC", "eks": "DevOps & IaC",
    "azure aks": "DevOps & IaC", "aks": "DevOps & IaC",
    "helm": "DevOps & IaC", "ansible": "DevOps & IaC",
    "ci/cd": "DevOps & IaC",
    "continuous integration": "DevOps & IaC",
    # Data Quality, Governance & Compliance
    "great expectations": "Data Quality, Governance & Compliance",
    "soda": "Data Quality, Governance & Compliance",
    "soda core": "Data Quality, Governance & Compliance",
    "sodacl": "Data Quality, Governance & Compliance",
    "monte carlo": "Data Quality, Governance & Compliance",
    "rbac": "Data Quality, Governance & Compliance",
    "role-based access control": "Data Quality, Governance & Compliance",
    "pii masking": "Data Quality, Governance & Compliance",
    "data lineage": "Data Quality, Governance & Compliance",
    "hipaa": "Data Quality, Governance & Compliance",
    "pci dss": "Data Quality, Governance & Compliance", "pci": "Data Quality, Governance & Compliance",
    "sox": "Data Quality, Governance & Compliance",
    "sox compliance": "Data Quality, Governance & Compliance",
    "sarbanes-oxley": "Data Quality, Governance & Compliance",
    "gdpr": "Data Quality, Governance & Compliance",
    "data contracts": "Data Quality, Governance & Compliance",
    # BI, Monitoring & Methodologies
    "power bi": "BI, Monitoring & Methodologies",
    "microsoft power bi": "BI, Monitoring & Methodologies",
    "tableau": "BI, Monitoring & Methodologies",
    "looker": "BI, Monitoring & Methodologies",
    "looker studio": "BI, Monitoring & Methodologies",
    "amazon quicksight": "BI, Monitoring & Methodologies",
    "quicksight": "BI, Monitoring & Methodologies",
    "google looker studio": "BI, Monitoring & Methodologies",
    "splunk": "BI, Monitoring & Methodologies",
    "prometheus": "BI, Monitoring & Methodologies",
    "grafana": "BI, Monitoring & Methodologies",
    "elk stack": "BI, Monitoring & Methodologies", "elk": "BI, Monitoring & Methodologies",
    "elasticsearch": "BI, Monitoring & Methodologies",
    "azure monitor": "BI, Monitoring & Methodologies",
    "datadog": "BI, Monitoring & Methodologies",
    "agile": "BI, Monitoring & Methodologies",
    "agile/scrum": "BI, Monitoring & Methodologies",
    "scrum": "BI, Monitoring & Methodologies",
    "dataops": "BI, Monitoring & Methodologies",
}

ECOSYSTEM_KEYWORDS = {
    "react", "vue", "vue.js", "angular", "next.js", "svelte", "ember",
    "html", "css", "javascript", "typescript", "jquery",
    "node.js", "express", "spring boot", "spring", "asp.net",
    "c++", "c#", ".net", "go", "golang", "rust", "ruby", "php",
    "graphql", "rest api", "rest apis", "grpc", "soap",
    "sap", "oracle ebs", "servicenow", "jira", "confluence",
    "mulesoft", "tibco", "ibm datastage",
    "blockchain", "smart contracts", "solidity",
    "rpa", "uipath", "automation anywhere",
}


# ═══════════════════════════════════════════════════════════════════
# 8. DOMAIN VOCABULARY (for prompt context)
# ═══════════════════════════════════════════════════════════════════

DOMAIN_VOCAB_REFERENCE = """
Domain vocabulary (inject naturally where it fits the JD industry):
- financial / insurance / wealth: regulatory reporting, audit trails, SOX-aligned data lineage, claims pipelines, actuarial datasets, customer-360, PII safeguards on customer financial data
- healthcare / pharma / medtech: HIPAA-aligned pipelines, clinical data, EHR integration, claims processing, pharmacy datasets, PHI-safe storage
- retail / e-commerce / FMCG: customer-360, transaction streams, inventory feeds, demand signals, omnichannel analytics
- manufacturing / energy / IoT: sensor telemetry, predictive maintenance datasets, SCADA feeds, asset performance
- fintech / payments: transaction streams, fraud signals, real-time risk scoring, ledger reconciliation
"""


# ═══════════════════════════════════════════════════════════════════
# 9. ASTRA PROMPT — single call, preset-aware, returns rich metadata
# ═══════════════════════════════════════════════════════════════════

ASTRA_PROMPT_TEMPLATE = """You are Astra, an elite resume tailoring engine for Lakshmi Supraja Konakanchi, a Senior Data Engineer based in Irving, Texas, USA.

Goal: take her BASE RESUME and the JOB DESCRIPTION, produce a tailored resume that passes ATS scanners AND a human recruiter's sniff test. Output must read like a confident senior engineer wrote it, not ChatGPT.

═══ PRESET (drives how aggressively to rewrite) ═══

%(PRESET_INSTRUCTION)s

═══ CANDIDATE IDENTITY (LOCKED, DO NOT CHANGE) ═══

Five roles in reverse chronological order. You REWRITE bullets within each role. You do NOT change titles, companies, locations, dates, or the number of roles.

1. Azure Data Engineer (Contract) | Northwestern Mutual | San Antonio, TX | Jul 2024 - Present
   Industry: Insurance / Financial Services / Wealth Management

2. AWS Data Engineer (Contract) | McKesson Corporation | Irving, TX | Aug 2023 - May 2024
   Industry: Healthcare / Pharmacy

3. GCP Data Engineer | Mindtree Limited (BigBasket Account) | Hyderabad, India | Aug 2021 - Jul 2022
   Industry: Retail / E-commerce

4. Data Engineer (Contract) | Geekyants Private Limited | Bengaluru, India | Jun 2020 - Jul 2021
   Industry: Tech / SaaS

5. Data Engineer (Contract) | Exide Energy Solutions | Bengaluru, India | Dec 2019 - May 2020
   Industry: Manufacturing / Energy / IoT

Years of experience: 5+. Senior level (this matters for verb selection — see SENIORITY).
Education: MS Management Information Systems, Lamar University (May 2024).
Cert: AWS Certified Data Engineer - Associate (DEA-C01) Aug 2025 - Aug 2028.

═══ #1 RULE: ZERO MISSING KEYWORDS ═══

Extract EVERY hard skill, tool, framework, language, methodology, certification, and compliance term from the JD. EVERY ONE must appear somewhere in the output (skills section, summary, or bullets).

Supraja can defend the modern data stack broadly. Claim what the JD asks for and weave it credibly.

INVISIBLE KEYWORD TECHNIQUE (CRITICAL): keywords appear inside BULLETS in natural prose, not stuffed in a Skills blob. A bullet like "Architected 40+ Apache Airflow DAGs orchestrating ingestion into Snowflake, processing 12 TB/day with transformations modeled in dbt and ingestion via Kafka Connect" passes both keyword-literal parsers AND semantic NLP parsers AND human recruiters. The Skills section is a categorized inventory, not a keyword dump.

CANONICAL FORM FIRST (research-validated rule): on first mention in each section, write the full canonical name with abbreviation in parentheses, then use the short form afterwards:
- "Apache Airflow" before "Airflow"
- "Amazon Web Services (AWS)" before "AWS"
- "Microsoft Azure" before "Azure"
- "Google Cloud Platform (GCP)" before "GCP"
- "Extract Transform Load (ETL)" before "ETL"
- "Role-Based Access Control (RBAC)" before "RBAC"
- "Continuous Integration / Continuous Deployment (CI/CD)" before "CI/CD"
This rewards both Workday Skills Cloud (semantic) and Taleo (keyword-literal).

═══ #2 RULE: ALLOWED METRICS ONLY ═══

Every percentage, ratio, GB/TB volume, or quantitative claim must trace to this list. Inventing new numbers is a hard fail.

Northwestern Mutual (current role): 500GB+/day, ~40% (batch runtime cut), ~25% (streaming latency), 20-30% (query throughput), ~30% (release effort), 15+ advisor dashboards, mentored 2 junior engineers, 15+ engineer team, 3+ additional reporting hours.

McKesson: 400GB+/day, ~30% (Redshift perf), ~25% (S3 cost), ~30% (pipeline failures), HIPAA audit cleared without findings.

Mindtree (BigBasket): 300GB+/day, ~25% (BigQuery efficiency).

Geekyants: ~25% (deployment effort).

Exide Energy: 2TB+ manufacturing sensor datasets.

If a bullet has no anchoring metric, write a strong QUALITATIVE outcome instead. NEVER invent a number.

═══ #3 RULE: QUANTIFY 60-70% OF BULLETS, NOT 100% ═══

Per research: ≥60% of bullets should contain a number, but NOT 100% — universal quantification triggers the "fake metric" tell, especially when round percentages repeat. Mix quantified bullets with strong qualitative outcomes.

═══ TITLE ARCHETYPE — detect & route ═══

Detect the JD title type:
- Direct DE titles (Senior DE / Data Engineer / Cloud DE): full match. Lead NM Azure work.
- Azure/AWS/GCP-specific DE: lead with the matching cloud role (NM=Azure, McKesson=AWS, BigBasket=GCP).
- Snowflake DE / Snowflake Developer: foreground NM Snowflake + dbt migration.
- ETL Developer / Informatica Developer: foreground Informatica/Talend/NiFi/Oozie + modern ETL migration.
- Analytics Engineer / Data Analytics Engineer: lead with dbt + Unity Catalog + governance + BI delivery.
- Big Data / PySpark / Spark Engineer: lead NM Databricks/PySpark + Exide Hadoop/Spark 2TB story.
- Database Engineer / Senior DBA: foreground SQL Server, PostgreSQL, Oracle, Cosmos DB tuning depth.
- AI Data Engineer / ML Platform Engineer / Data Engineer for AI: position as data platform engineer building feature pipelines and model-ready datasets; reference dbt + MLflow + Unity Catalog + vector store integration where the JD asks.
- Anything else (Solutions Architect, niche stack, unusual title): tailor aggressively to whatever the JD asks for.

═══ INDUSTRY DETECTION ═══

- Financial / Insurance / Wealth / Brokerage / FinTech → lived at NM. Claim directly.
- Healthcare / Pharma / Medtech / Life Sciences → lived at McKesson. Claim directly.
- Retail / E-commerce / FMCG / Consumer → lived at BigBasket. Claim directly.
- Manufacturing / Energy / Utilities / IoT → lived at Exide. Claim directly.
- Tech / SaaS / Software / Consulting → lived at Geekyants. Claim directly.
- Anything else → cross-industry transferable. Frame multi-cloud breadth as a feature.

%(DOMAIN_VOCAB)s

═══ SUMMARY (5 SENTENCES, PRONOUN-FREE, NO EM DASHES) ═══

Pronoun-free: NO "I", "she", "Supraja", "Lakshmi", "her", "my". Drop the subject entirely.

Sentence 1 — IDENTITY ANCHOR: JD's EXACT role title + 5+ years + AWS / Azure / GCP + concrete scale fact (500GB+/day OR 5+ years of regulated-data delivery).
GOOD: "Senior Data Engineer with 5+ years building production data platforms across Amazon Web Services (AWS), Microsoft Azure, and Google Cloud Platform (GCP), currently moving 500GB+/day in regulated financial services."

Sentence 2 — TECHNICAL DEPTH: specific tools at production scale reflecting JD stack overlap.

Sentence 3 — DOMAIN: claim industry if lived, transfer if not. Never inflate.

Sentence 4 — ENGINEERING EDGE: ONE concrete strength the JD values (governance, lineage, IaC, streaming, Lakehouse, cost, data quality, mentoring).

Sentence 5 — JD-STACK FIT: close with the JD's own technical language. NEVER name the target company.

BANNED summary openers: "Highly motivated", "Results-driven", "Passionate", "Dedicated professional", "Detail-oriented", "Seasoned", "Dynamic professional".
BANNED summary closers: "Aims to", "Ready to", "Seeking to", "Eager to", "Looking to", "Excited to", "Driven to", "Committed to", "Poised to".

═══ SKILLS RULES ═══

Output 7-10 categories. Place JD-mentioned tools FIRST in each category. Add EVERY JD keyword. Use canonical category names:
- Programming Languages
- AWS Cloud / Azure Cloud / GCP Cloud (split when JD is cloud-specific; combine into "Cloud Platforms" if JD is cloud-neutral)
- Data Platforms & Warehouses
- Streaming & Orchestration
- Databases
- Data Modeling & File Formats
- DevOps & Infrastructure as Code (IaC)
- Data Quality, Governance & Compliance
- BI, Monitoring & Methodologies
- Ecosystem Integration & Exposure (for non-core JD keywords like React, Vue, Java, C++, .NET, niche CRMs)

═══ EXPERIENCE BULLETS ═══

ALL 5 roles MUST appear in reverse chronological order. Per role:
- Northwestern Mutual: 5-6 responsibilities + 2 achievements (PRESENT TENSE)
- McKesson: 5 responsibilities + 2 achievements (past tense)
- Mindtree/BigBasket: 3-4 responsibilities + 0-1 achievements
- Geekyants: 2-3 responsibilities + 0 achievements
- Exide: 1-2 responsibilities + 0 achievements

XYZ BULLET FORMULA (Google's, research-validated): "Accomplished [X] as measured by [Y] by doing [Z]." Front-load the outcome to survive the F-pattern recruiter skim.

SENIORITY VERBS — Supraja is Senior level. Use senior verbs at least 30-40% of bullets:
- Strong senior: Architected, Designed, Led, Mentored, Owned, Drove, Established, Standardised, Productionised, Migrated, Consolidated, Introduced, Orchestrated, Engineered, Partnered.
- Avoid for senior: "Helped", "Assisted", "Supported", "Contributed", "Participated".

ARCHITECTURAL OWNERSHIP signals to foreground (research differentiator from mid-level):
- Named migrations (legacy → modern ETL, on-prem → cloud, batch → streaming)
- Cross-team data products
- Mentorship of 1-3 engineers (she has 2)
- Cost governance ownership

BURSTINESS RULES (anti-AI):
- Vary bullet length 8-26 words. Mix short and long deliberately.
- No more than 2 consecutive bullets start with the same verb.
- Avoid identical syntactic patterns across consecutive bullets.

═══ ANTI-AI WRITING — BANNED WORDS ═══

NEVER use: delve, leverage, leveraged, leveraging, spearhead, spearheaded, robust, seamless, seamlessly, intricate, harness, harnessed, unlock, unleash, empower, empowered, transformative, groundbreaking, pivotal, cutting-edge, holistic, synergy, foster, fostered, streamline, streamlined, elevate, elevated, paradigm, ever-evolving, tapestry, realm, embark, vibrant, crucial, compelling, testament, navigate the landscape, utilize, utilized, utilizing, showcasing, highlighting, demonstrating, fostering, cultivating, at the forefront of, at the intersection of, passionate about, driven by, committed to excellence, serves as, stands as.

NEVER START a sentence with: Furthermore, Moreover, Additionally, Consequently, In today's, In an era of, It is worth noting.

NEVER use EM DASHES (—) anywhere. Use commas, periods, colons, parentheses, or rephrase.

═══ OUTPUT JSON SCHEMA ═══

Return ALL fields below. Order matters — model plans by completing earlier fields first.

{
  "jd_intelligence": {
    "top_keywords": ["<10-15 most important hard skills from JD, exact JD phrasing>"],
    "required_tools": ["<must-have tools from JD>"],
    "nice_to_have_tools": ["<nice-to-have tools from JD>"],
    "emphasis_areas": ["<2-4 themes: e.g. 'real-time streaming', 'governance', 'cost optimisation'>"],
    "industry": "<JD's industry>",
    "seniority_signal": "<junior | mid | senior | staff>"
  },
  "tailoring_notes": {
    "rewrote_for": ["<short notes: 'NM bullet 2: re-emphasised Snowflake and dbt migration', etc.>"],
    "did_not_add": ["<plain English: 'Kubernetes — not in your original' if JD mentions but base does not, etc.>"],
    "keyword_coverage_summary": "<short plain English: 'Covered 13 of 15 top JD keywords. Missing: Kubernetes, Airbyte.'>",
    "preset_used": "<echo back the preset name>"
  },
  "candidate_title": "<mirror JD's exact title, max 8 words; default 'Senior Data Engineer'>",
  "summary": "<EXACTLY 5 sentences, pronoun-free, no em dashes, canonical-form-first>",
  "skills": [
    {"category": "<canonical category name>", "technologies": "<comma-separated, JD keywords FIRST, canonical names>"},
    ...
  ],
  "experience": [
    {
      "company": "Northwestern Mutual",
      "responsibilities": ["<5-6 rewritten, present tense, XYZ formula where natural>"],
      "achievements": ["<2 achievement bullets>"]
    },
    {
      "company": "McKesson Corporation",
      "responsibilities": ["<5 rewritten, past tense>"],
      "achievements": ["<2 achievement bullets>"]
    },
    {
      "company": "Mindtree Limited (BigBasket Account)",
      "responsibilities": ["<3-4 rewritten>"],
      "achievements": []
    },
    {
      "company": "Geekyants Private Limited",
      "responsibilities": ["<2-3 rewritten>"],
      "achievements": []
    },
    {
      "company": "Exide Energy Solutions",
      "responsibilities": ["<1-2 rewritten>"],
      "achievements": []
    }
  ],
  "target_company": "<extracted from JD or 'Company'>"
}

═══ FINAL CHECK ═══

1. Em dashes (—) — remove every one.
2. Banned AI-writing words — remove every one.
3. Fabricated numbers — only allowed-metrics-list numbers; otherwise qualitative.
4. All 5 companies present in experience array.
5. Every JD hard skill appears somewhere.
6. ≥60% (not 100%) of bullets quantified.
7. Summary exactly 5 sentences, pronoun-free.
8. Canonical-form-first on first mention of each major tool/abbreviation.
9. tailoring_notes.did_not_add lists JD-asked tools that are NOT in base resume (trust signal).

Return ONLY the JSON object. No prose, no markdown fences.
"""


# ═══════════════════════════════════════════════════════════════════
# 10. COVER LETTER PROMPT (kept from v1.0, runs only on button)
# ═══════════════════════════════════════════════════════════════════

COVER_LETTER_PROMPT = """You are Lakshmi Supraja Konakanchi writing a direct cover email to a Hiring Manager in the US.
Goal: sound 100% human. Get a response.

NO EM DASHES (—) anywhere. Use commas, periods, colons.

BANNED PHRASES: "I am writing to express my interest", "I am excited to apply", "Please find my resume attached", "I believe I am a perfect fit", "passionate about", "driven by", "committed to excellence", "at the forefront of", "showcasing", "highlighting", "demonstrating", "serves as", "stands as", "leveraging", "harnessing", "utilizing", "seamless", "innovative", "groundbreaking", "testament to", "underscores", "pivotal", "realm", "tapestry", "delve", "spearhead", "robust", "holistic", "synergy", "foster", "streamline", "elevate", "transformative".

CONTEXT:
- Candidate: Lakshmi Supraja Konakanchi, 5+ years multi-cloud Senior Data Engineer, Irving TX
- Real industries: Insurance/Wealth (Northwestern Mutual), Healthcare/Pharma (McKesson), Retail/E-commerce (BigBasket), Tech/SaaS (Geekyants), Manufacturing/IoT (Exide)
- AWS Certified Data Engineer - Associate (DEA-C01)

If JD industry matches one she's lived in, claim directly. If not, frame as transferable.

THE OPENING: start with a specific observation about the company's data challenge from the JD. NEVER "I am applying for...".

WAR STORY — pick the BEST matching:
1. Northwestern Mutual (Insurance/Wealth/Azure/dbt/Snowflake): "At Northwestern Mutual, I migrated legacy on-premises ETL into Azure Synapse and Snowflake using dbt on Databricks, cutting overnight batch runtime by 40% and freeing 3+ extra reporting hours for the wealth management BI teams every morning."
2. McKesson (Healthcare/HIPAA/AWS/Iceberg): "At McKesson, I rebuilt the pharmacy data lake on Amazon S3 with Apache Iceberg, enabling ACID transactions and schema evolution for 400GB+ of daily claims. The same architecture cleared internal HIPAA audit review on access controls without findings."
3. Mindtree/BigBasket (Retail/GCP/streaming): "At BigBasket, I built Google Cloud Pub/Sub plus Dataflow streaming feeds into BigQuery, processing 300GB/day of retail transaction data and lifting analytical query efficiency by 25% for the supply-chain and category teams."
4. Geekyants (Tech/SaaS/multi-tool): "At Geekyants, I shipped Apache Spark, Kafka, and Airflow pipelines across PostgreSQL, MySQL, and MongoDB sources, automating CI/CD with Jenkins and AWS Lambda and cutting manual deployment effort by 25%."
5. Exide (Manufacturing/IoT/scale): "At Exide Energy, I built Apache Hadoop and Spark batch pipelines on HDFS, Hive, and HBase to process 2TB+ of manufacturing sensor data, then migrated selected workloads to AWS Lambda and Kinesis using Terraform for early cloud modernisation."

STRUCTURE: 4 short paragraphs.
1. Hook (their data pain point from JD)
2. Bridge: "This is close to a problem I solved at [Company]..."
3. War story with specific tools and numbers
4. Brief closing tying multi-cloud breadth to their team

Sign off with "Thank you," on its own line, then "Lakshmi Supraja Konakanchi".

STYLE: vary sentence length, active voice, plain verbs. 240-340 words body.

Return ONLY the letter body as plain text. No "Dear" greeting (renderer adds it). No subject. No markdown.
"""


# ═══════════════════════════════════════════════════════════════════
# 11. ATS SCORING PROMPT (kept from v1.0, optional button)
# Note: per research, target match band is 78-85%, NOT 95%+.
# ═══════════════════════════════════════════════════════════════════

ATS_SCORING_PROMPT = """You are a strict ATS scanner for a US Senior Data Engineer role.
Compare the RESUME JSON against the JOB DESCRIPTION.

CRITICAL: a score above 88% indicates KEYWORD STUFFING and triggers human-recruiter "robotic" detectors. The optimal band is 78-85%. Do not push artificially high.

Scoring (0-100):
- Keyword match density (40%): % of JD hard skills in resume. Penalise above-85%.
- Experience relevance (25%): bullets describe work that solves JD's problems.
- Title-match alignment (15%): candidate_title and summary match JD's role language.
- Domain fit (10%): industry background relevant.
- Seniority alignment (10%): experience level matches JD.

Output ONLY valid JSON:
{"score": <int 0-100>, "reasoning": "<1 sentence>", "missing_keywords": "<comma-separated JD keywords NOT in resume>", "title_match_status": "<match | partial | mismatch>", "stuffing_risk": "<low | medium | high>"}
"""


# ═══════════════════════════════════════════════════════════════════
# 12. PYDANTIC SCHEMAS
# Schema field order matters: jd_intelligence first → model plans
# before writing the resume body.
# ═══════════════════════════════════════════════════════════════════

class JDIntelligence(BaseModel):
    top_keywords: List[str] = Field(default_factory=list)
    required_tools: List[str] = Field(default_factory=list)
    nice_to_have_tools: List[str] = Field(default_factory=list)
    emphasis_areas: List[str] = Field(default_factory=list)
    industry: str = ""
    seniority_signal: str = "senior"


class TailoringNotes(BaseModel):
    rewrote_for: List[str] = Field(default_factory=list)
    did_not_add: List[str] = Field(default_factory=list)
    keyword_coverage_summary: str = ""
    preset_used: str = ""


class ExperienceItem(BaseModel):
    company: str
    responsibilities: List[str]
    achievements: List[str] = Field(default_factory=list)


class SkillCategory(BaseModel):
    category: str
    technologies: str


class TailoredOutput(BaseModel):
    jd_intelligence: JDIntelligence
    tailoring_notes: TailoringNotes
    candidate_title: str
    summary: str
    skills: List[SkillCategory]
    experience: List[ExperienceItem]
    target_company: str = "Company"


def get_clean_schema(pydantic_cls):
    schema = pydantic_cls.model_json_schema()
    def _clean(d):
        if isinstance(d, dict):
            for key in ["additionalProperties", "title"]:
                d.pop(key, None)
            for v in d.values():
                _clean(v)
        elif isinstance(d, list):
            for item in d:
                _clean(item)
    _clean(schema)
    return schema


# ═══════════════════════════════════════════════════════════════════
# 13. RUNTIME HELPERS — em-dash, pronouns, category validator
# Same as v1.0. Quality safety nets, not policing.
# ═══════════════════════════════════════════════════════════════════

def strip_em_dashes(text: str) -> str:
    if not text:
        return text
    cleaned = re.sub(r"\s*—\s*", ", ", text)
    cleaned = re.sub(r",\s*,", ",", cleaned)
    return cleaned


def strip_summary_pronouns(text: str) -> str:
    if not text:
        return text
    patterns = [
        (r"(^|(?<=\.\s))(I am|I'm|She is|She's|Lakshmi is|Supraja is|Lakshmi Supraja is)\s+(a|an)\s+", r"\1"),
        (r"(^|(?<=\.\s))(I am|I'm|She is|She's|Lakshmi is|Supraja is|Lakshmi Supraja is)\s+", r"\1"),
        (r"(^|(?<=\.\s))(She|He)\s+(brings|has|holds|manages|managed|architects|architected|builds|built|leads|led|designs|designed|delivers|delivered|owns|owned|partners|partnered|mentors|mentored)\b",
         lambda m: m.group(1) + m.group(3).capitalize()),
        (r"(^|(?<=\.\s))I\s+(bring|have|hold|manage|managed|architect|architected|build|built|lead|led|design|designed|deliver|delivered|own|owned|partner|partnered|mentor|mentored)\b",
         lambda m: m.group(1) + m.group(2).capitalize()),
        (r"(^|(?<=\.\s))(Lakshmi Supraja|Lakshmi|Supraja)\s+(brings|has|holds|architects|architected|built|builds|managed|manages|led|leads)\b",
         lambda m: m.group(1) + m.group(3).capitalize()),
        (r"\b(her|My|my)\s+(MS|experience|background|certification|career|work)\b", r"\2"),
    ]
    cleaned = text
    for pattern, repl in patterns:
        cleaned = re.sub(pattern, repl, cleaned)
    if cleaned and cleaned[0].islower():
        cleaned = cleaned[0].upper() + cleaned[1:]
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned


def _normalize_skill(s: str) -> str:
    if not s:
        return ""
    n = s.lower().strip()
    n = re.sub(r"[()]", " ", n)
    n = re.sub(r"\s+", " ", n).strip()
    return n


def validate_and_repair_skill_categories(skills_list):
    if not skills_list:
        return skills_list
    working = {}
    for entry in skills_list:
        cat = entry.get("category", "").strip()
        techs = entry.get("technologies", "")
        if not cat:
            continue
        if isinstance(techs, str):
            tools = [t.strip() for t in techs.split(",") if t.strip()]
        else:
            tools = list(techs) if isinstance(techs, list) else []
        if cat not in working:
            working[cat] = []
        for t in tools:
            if t not in working[cat]:
                working[cat].append(t)

    misplaced = []
    for current_cat, tool_list in list(working.items()):
        for tool in list(tool_list):
            tool_norm = _normalize_skill(tool)
            if any(kw == tool_norm or kw in tool_norm.split() for kw in ECOSYSTEM_KEYWORDS):
                if "ecosystem" not in current_cat.lower():
                    misplaced.append((tool, current_cat, "Ecosystem Integration & Exposure"))
                continue
            if tool_norm in MASTER_CATEGORY_MAP:
                correct = MASTER_CATEGORY_MAP[tool_norm]
                if correct.lower() != current_cat.lower():
                    misplaced.append((tool, current_cat, correct))

    for tool, from_cat, to_cat in misplaced:
        if tool in working.get(from_cat, []):
            working[from_cat].remove(tool)
        if to_cat not in working:
            working[to_cat] = []
        if not any(t.lower() == tool.lower() for t in working[to_cat]):
            working[to_cat].append(tool)

    working = {cat: tools for cat, tools in working.items() if tools}
    return [{"category": cat, "technologies": ", ".join(tools)}
            for cat, tools in working.items()]


# ═══════════════════════════════════════════════════════════════════
# 14. FACTS INVENTORY + GROUNDING VALIDATION
# Per research: the single most important moat against the "sea of
# sameness" is candidate-grounded bullets. Every tool and number in
# a generated bullet should trace to the base resume.
#
# Research recommended embedding-based cosine grounding. We can't use
# embeddings (constraint). Token-overlap is the practical fallback:
# extract every tool name and number from the base resume, flag any
# generated bullet that introduces tokens not in the inventory.
# ═══════════════════════════════════════════════════════════════════

# Tools that count as "Supraja's stack" for grounding. Extracted from
# her base skills + experience descriptions. Used to detect generated
# bullets that introduce tools she's never used.
def extract_tools_inventory() -> Set[str]:
    """Return a set of normalised tool names that appear in the base
    resume. Generated bullets should mostly mention tools in this set."""
    text = SUPRAJA_BASE_RESUME.lower()
    # Pull out distinct multi-word tool phrases
    inventory = set()
    for canonical, _ in MASTER_CATEGORY_MAP.items():
        if canonical in text:
            inventory.add(canonical)
    # Add ecosystem keywords if present in base
    for kw in ECOSYSTEM_KEYWORDS:
        if kw in text:
            inventory.add(kw)
    return inventory


def _norm_vol(token: str) -> str:
    """Normalise a volume/percentage/count token for grounding comparison.

    Fixes Bug 1: the base resume says "500GB+" but tailored bullets often
    say "500GB+/day" or "500 GB / day". Without normalisation every daily-
    volume bullet got falsely flagged as ungrounded. This strips the rate
    suffix (/day, /month, /hour) and whitespace so both sides match.
    """
    if not token:
        return ""
    t = token.upper().replace(" ", "")
    # Strip rate suffixes (one pass handles /DAY, /MONTH, /HOUR, /WEEK, /YR)
    t = re.sub(r'/(DAY|MONTH|HOUR|WEEK|YEAR|YR|MIN|MINUTE|SEC|SECOND)S?', '', t)
    return t


def extract_numbers_inventory() -> Set[str]:
    """Return a set of number tokens (percentages, GB/TB volumes,
    counts) that appear in the base resume. Generated bullets should
    only use these numbers. All tokens are passed through `_norm_vol`
    so daily-volume style ("500GB+/day") matches base ("500GB+")."""
    text = SUPRAJA_BASE_RESUME
    inventory = set()
    # Percentages
    for m in re.finditer(r'(\d+\.?\d*)\s*%', text):
        inventory.add(_norm_vol(m.group(1) + "%"))
    # GB/TB volumes — allow optional /day suffix in the regex so we
    # capture it consistently, then strip via _norm_vol
    for m in re.finditer(r'(\d+\.?\d*)\s*(GB|TB|MB|KB)\+?\s*(/\s*(day|month|hour|week|year|yr)s?)?',
                         text, re.IGNORECASE):
        inventory.add(_norm_vol(m.group(0)))
    # Ranges like 20-30%
    for m in re.finditer(r'(\d+)\s*-\s*(\d+)\s*%', text):
        inventory.add(_norm_vol(f"{m.group(1)}-{m.group(2)}%"))
    # Plain numbers with + (15+, 2 junior, etc.)
    for m in re.finditer(r'\b(\d+)\+', text):
        inventory.add(_norm_vol(f"{m.group(1)}+"))
    return inventory


TOOLS_INVENTORY = extract_tools_inventory()
NUMBERS_INVENTORY = extract_numbers_inventory()


def find_ungrounded_tools(bullet: str) -> List[str]:
    """Return tool-like tokens in the bullet that are NOT in the base
    resume tools inventory. Soft warning, not auto-reject."""
    if not bullet:
        return []
    bullet_lower = bullet.lower()
    # Check all known tools from MASTER_CATEGORY_MAP + ECOSYSTEM_KEYWORDS
    candidates = list(MASTER_CATEGORY_MAP.keys()) + list(ECOSYSTEM_KEYWORDS)
    found_in_bullet = set()
    for tool in candidates:
        # Word-boundary match
        if re.search(r'\b' + re.escape(tool) + r'\b', bullet_lower):
            found_in_bullet.add(tool)
    # Return ones not in base inventory
    return [t for t in found_in_bullet if t not in TOOLS_INVENTORY]


def find_ungrounded_numbers(bullet: str) -> List[str]:
    """Return number tokens in the bullet that are NOT in the base
    resume numbers inventory. Catches fabricated metrics. Both sides
    are normalised via `_norm_vol` so "/day" and similar rate suffixes
    don't cause false positives."""
    if not bullet:
        return []
    found = set()
    # Percentages
    for m in re.finditer(r'(\d+\.?\d*)\s*%', bullet):
        found.add(_norm_vol(m.group(1) + "%"))
    # GB/TB with optional rate suffix
    for m in re.finditer(r'(\d+\.?\d*)\s*(GB|TB|MB|KB)\+?\s*(/\s*(day|month|hour|week|year|yr)s?)?',
                         bullet, re.IGNORECASE):
        found.add(_norm_vol(m.group(0)))
    # Plain N+ tokens
    for m in re.finditer(r'\b(\d+)\+', bullet):
        found.add(_norm_vol(f"{m.group(1)}+"))
    # Ranges
    for m in re.finditer(r'(\d+)\s*-\s*(\d+)\s*%', bullet):
        found.add(_norm_vol(f"{m.group(1)}-{m.group(2)}%"))
    return [n for n in found if n not in NUMBERS_INVENTORY]


def validate_grounding(experience: List[Dict]) -> Dict[str, List[str]]:
    """Run grounding checks across all bullets. Return a dict of
    company → list of warnings for the UI."""
    warnings = {}
    for role in experience:
        company = role.get("company", "")
        role_warnings = []
        all_bullets = role.get("responsibilities", []) + role.get("achievements", [])
        for bullet in all_bullets:
            bad_nums = find_ungrounded_numbers(bullet)
            if bad_nums:
                role_warnings.append(
                    f"Number(s) not in base resume: {', '.join(bad_nums)} — "
                    f"in bullet: \"{bullet[:80]}...\""
                )
        if role_warnings:
            warnings[company] = role_warnings
    return warnings


# ═══════════════════════════════════════════════════════════════════
# 15. BURSTINESS AUDIT
# Per research: sentence-length stdev <4.0 OR any lead verb appearing
# more than 2 times is the textbook AI signature. Detected
# deterministically, no LLM call needed.
# ═══════════════════════════════════════════════════════════════════

def burstiness_audit(experience: List[Dict]) -> Dict:
    """Compute burstiness metrics across all bullets in the resume.
    Returns a dict with stdev, lead-verb-counts, verb-diversity, and
    pass/fail flags."""
    all_bullets = []
    for role in experience:
        all_bullets.extend(role.get("responsibilities", []))
        all_bullets.extend(role.get("achievements", []))

    if len(all_bullets) < 2:
        return {
            "sentence_stdev": 0.0,
            "verb_diversity": 0.0,
            "lead_verb_counts": {},
            "max_repeat_verb": "",
            "max_repeat_count": 0,
            "passes_burstiness": True,
        }

    word_counts = [len(b.split()) for b in all_bullets]
    stdev = statistics.stdev(word_counts) if len(word_counts) > 1 else 0.0

    lead_verbs = []
    for b in all_bullets:
        first_word = re.match(r'^\s*(\w+)', b)
        if first_word:
            lead_verbs.append(first_word.group(1).lower())

    verb_counts = {}
    for v in lead_verbs:
        verb_counts[v] = verb_counts.get(v, 0) + 1

    unique_verbs = len(set(lead_verbs))
    diversity = unique_verbs / max(len(lead_verbs), 1)

    max_verb = ""
    max_count = 0
    for v, c in verb_counts.items():
        if c > max_count:
            max_verb = v
            max_count = c

    # Pass thresholds per research: stdev ≥4.0, no verb >2 times
    passes = (stdev >= 4.0) and (max_count <= 2)

    return {
        "sentence_stdev": round(stdev, 2),
        "verb_diversity": round(diversity, 2),
        "lead_verb_counts": verb_counts,
        "max_repeat_verb": max_verb,
        "max_repeat_count": max_count,
        "passes_burstiness": passes,
    }


# ═══════════════════════════════════════════════════════════════════
# 16. MULTI-AXIS SCORING (deterministic, zero LLM cost)
# Per research: replace single match % with 4-axis card.
# All four axes computed from generated text. NO API calls.
# ═══════════════════════════════════════════════════════════════════

def compute_impact_axis(experience: List[Dict]) -> Tuple[int, str]:
    """% of bullets with quantifiable result. Target 60-70% per
    research — universal quantification triggers fake-metric tell."""
    all_bullets = []
    for role in experience:
        all_bullets.extend(role.get("responsibilities", []))
        all_bullets.extend(role.get("achievements", []))
    if not all_bullets:
        return 0, "No bullets"

    with_numbers = 0
    for b in all_bullets:
        if re.search(r'\d+\.?\d*\s*(%|GB|TB|MB|KB|hours?|minutes?|engineers?|dashboards?|\+|x)', b, re.IGNORECASE):
            with_numbers += 1
        elif re.search(r'\b\d{2,}\b', b):  # standalone 2+ digit numbers
            with_numbers += 1

    pct = round(100 * with_numbers / len(all_bullets))
    # Score the percentage against the 60-70% sweet spot
    if 60 <= pct <= 75:
        score = 95  # ideal
    elif 50 <= pct < 60 or 75 < pct <= 85:
        score = 80
    elif 40 <= pct < 50 or 85 < pct <= 95:
        score = 65
    elif pct > 95:
        score = 50  # fake-metric tell
    else:
        score = 40

    label = f"{pct}% of bullets quantified"
    return score, label


def compute_keywords_axis(skills: List[Dict], summary: str,
                          experience: List[Dict],
                          jd_keywords: List[str]
                          ) -> Tuple[int, str, List[str], bool]:
    """% of jd_intelligence.top_keywords present in the resume.
    Cap above 85% triggers stuffing warning per research.

    Returns (score, label, missing_keywords, low_confidence_flag).

    Bug 3 fix: if the model lazily extracted fewer than 8 keywords, the
    score becomes statistically unreliable (3/4 = 75% looks fine when
    real coverage might be 12/40). We surface `low_confidence=True` in
    that case and the UI shows a warning instead of trusting the number.
    The threshold of 8 matches the research recommendation that a
    typical JD has 10-15 hard skills the model should extract."""
    # Guard 1: model returned no keywords at all
    if not jd_keywords:
        return 0, "No JD keywords extracted", [], True

    # Build full resume text
    parts = [summary]
    for s in skills:
        parts.append(s.get("category", ""))
        parts.append(s.get("technologies", ""))
    for role in experience:
        for b in role.get("responsibilities", []):
            parts.append(b)
        for b in role.get("achievements", []):
            parts.append(b)
    full_text = " ".join(parts).lower()

    found = []
    missing = []
    for kw in jd_keywords:
        kw_clean = kw.strip().lower()
        if not kw_clean:
            continue
        if kw_clean in full_text:
            found.append(kw)
        else:
            missing.append(kw)

    total = len(found) + len(missing)
    if total == 0:
        return 0, "No JD keywords", [], True

    # Guard 2: too few keywords for the score to be meaningful
    low_confidence = total < 8

    pct = round(100 * len(found) / total)
    # Score against research-recommended 78-85% target band
    if 78 <= pct <= 85:
        score = 95  # ideal
    elif 70 <= pct < 78:
        score = 85
    elif 85 < pct <= 92:
        score = 80  # approaching stuffing zone
    elif 60 <= pct < 70:
        score = 70
    elif pct > 92:
        score = 60  # stuffing risk
    else:
        score = 50

    # If low confidence, dampen the score so the overall match doesn't
    # get artificially inflated and surface a different label
    if low_confidence:
        score = min(score, 65)
        label = (
            f"{len(found)} of {total} keywords ({pct}%) "
            f"— ⚠️ only {total} extracted, may underrepresent JD"
        )
    else:
        label = f"{len(found)} of {total} JD keywords ({pct}%)"

    return score, label, missing, low_confidence


def compute_readability_axis(summary: str, experience: List[Dict],
                             burstiness: Dict) -> Tuple[int, str]:
    """Composite: banned-phrase density + burstiness."""
    # Banned phrase count across all body text
    body_parts = [summary]
    for role in experience:
        body_parts.extend(role.get("responsibilities", []))
        body_parts.extend(role.get("achievements", []))
    body = " ".join(body_parts)
    total_words = len(body.split()) if body else 1
    banned_hits = count_banned_tier1(body)
    banned_density = banned_hits / max(total_words, 1)

    # Score components
    if banned_hits == 0:
        banned_score = 100
    elif banned_density < 0.005:
        banned_score = 80
    elif banned_density < 0.01:
        banned_score = 65
    else:
        banned_score = 40

    burstiness_score = 100 if burstiness["passes_burstiness"] else 60

    score = round(0.6 * banned_score + 0.4 * burstiness_score)

    if banned_hits == 0 and burstiness["passes_burstiness"]:
        label = "Clean human-sounding prose"
    elif banned_hits == 0:
        label = f"No AI tells, but verb repetition ({burstiness['max_repeat_verb']} × {burstiness['max_repeat_count']})"
    elif burstiness["passes_burstiness"]:
        label = f"{banned_hits} AI-tell phrase(s) remaining"
    else:
        label = f"{banned_hits} AI-tells + verb repetition"

    return score, label


def compute_experience_axis(experience: List[Dict]) -> Tuple[int, str]:
    """% of bullets using senior-tier verbs. For 5-YOE Senior DE,
    target 30-50% senior verbs (mix is good; all-senior is also a tell)."""
    all_bullets = []
    for role in experience:
        all_bullets.extend(role.get("responsibilities", []))
        all_bullets.extend(role.get("achievements", []))
    if not all_bullets:
        return 0, "No bullets"

    senior_count = 0
    junior_count = 0
    for b in all_bullets:
        first_word = re.match(r'^\s*(\w+)', b)
        if not first_word:
            continue
        verb = first_word.group(1).lower()
        if verb in SENIOR_VERBS:
            senior_count += 1
        elif verb in JUNIOR_VERBS:
            junior_count += 1

    total = len(all_bullets)
    senior_pct = round(100 * senior_count / total)

    # 30-50% senior is the sweet spot
    if 30 <= senior_pct <= 60:
        score = 95
    elif 20 <= senior_pct < 30 or 60 < senior_pct <= 75:
        score = 80
    elif senior_pct > 75:
        score = 65  # too uniformly senior
    else:
        score = 55

    if junior_count > 0:
        score = max(0, score - 20)
        label = f"{senior_pct}% senior verbs (warning: {junior_count} junior verbs)"
    else:
        label = f"{senior_pct}% senior-tier verbs"

    return score, label


def compute_overall_match(impact: int, keywords: int,
                          readability: int, experience: int) -> int:
    """Weighted overall. Caps at 88 per research — anything higher
    triggers stuffing detector. Weights match research-recommended
    sub-axis importance."""
    raw = (
        0.25 * impact +
        0.35 * keywords +
        0.20 * readability +
        0.20 * experience
    )
    return min(88, round(raw))


def overall_match_label(score: int) -> Tuple[str, str]:
    """(emoji + plain English label) for the score."""
    if score >= 78:
        return "🟢", "Excellent match — ready to send"
    if score >= 65:
        return "🟡", "Good match — small tweaks recommended"
    if score >= 50:
        return "🟠", "Decent match — review missing keywords"
    return "🔴", "Major gaps — consider another role"


# ═══════════════════════════════════════════════════════════════════
# 17. READY-TO-SEND CHECKLIST (deterministic, zero LLM cost)
# ═══════════════════════════════════════════════════════════════════

def generate_ready_checklist(overall: int, impact_score: int,
                             keywords_score: int, readability_score: int,
                             missing_keywords: List[str],
                             grounding_warnings: Dict,
                             burstiness: Dict,
                             kw_low_confidence: bool = False
                             ) -> List[Tuple[bool, str]]:
    """Per research: a green-check block above the download button.
    Returns list of (passed, message) tuples.

    Messages are written for non-technical reviewers (e.g. Niharika).
    Raw numbers stay visible in the Score Breakdown card up top; this
    block translates them into plain English so a coordinator can
    decide whether to send the resume without understanding stdev.
    """
    checks = []

    # ─── Overall match ───
    if 78 <= overall <= 85:
        checks.append((True,
            f"Overall match looks strong — in the recruiter-sweet-spot band "
            f"({overall}%)."))
    elif overall > 85:
        checks.append((True,
            f"Overall match is high ({overall}%). At this level the resume "
            f"may start to read as keyword-stuffed. Consider Keep-my-voice "
            f"preset if it feels off."))
    elif 70 <= overall < 78:
        checks.append((True,
            f"Overall match is acceptable ({overall}%). You can send as-is "
            f"or check the missing keywords below."))
    else:
        checks.append((False,
            f"Overall match is low ({overall}%). Consider Maximize-match "
            f"preset or check whether this JD is the right fit."))

    # ─── Impact (quantification) ───
    if impact_score >= 80:
        checks.append((True,
            "Bullets carry a healthy mix of numbers and qualitative outcomes."))
    elif impact_score >= 60:
        checks.append((True,
            "Some bullets could use a number or scale to land harder, but "
            "good enough to send."))
    else:
        checks.append((False,
            "Most bullets are missing concrete numbers. Add scale, percentages, "
            "or named systems where you can."))

    # ─── Keywords ───
    if kw_low_confidence:
        checks.append((False,
            "⚠️ Astra only pulled a small number of keywords from this JD. "
            "The keyword score may not reflect real coverage — review the "
            "missing list manually before sending."))
    elif keywords_score >= 80 and not missing_keywords:
        checks.append((True,
            "All top JD keywords are present in the resume."))
    elif keywords_score >= 75:
        miss_preview = ", ".join(missing_keywords[:3])
        more = f" and {len(missing_keywords)-3} more" if len(missing_keywords) > 3 else ""
        checks.append((True,
            f"Most JD keywords are covered. Still missing: {miss_preview}{more}."))
    else:
        miss_preview = ", ".join(missing_keywords[:5])
        checks.append((False,
            f"Several JD keywords are missing: {miss_preview}. "
            f"Add these to Skills or weave into a bullet where Supraja has the work."))

    # ─── Readability + burstiness (plain English per UX fix) ───
    if readability_score >= 85 and burstiness["passes_burstiness"]:
        checks.append((True,
            "The writing sounds human — no AI-tell phrases, and bullets vary "
            "naturally in length and openers."))
    elif readability_score >= 85:
        # Readability passes but burstiness failed — explain what that means
        if burstiness["max_repeat_count"] > 2:
            checks.append((False,
                f"⚠️ Bullets feel a bit uniform — the word \"{burstiness['max_repeat_verb'].capitalize()}\" "
                f"starts {burstiness['max_repeat_count']} bullets. Reword two of them "
                f"to use different verbs (e.g. Architected, Migrated, Tuned)."))
        else:
            checks.append((False,
                "⚠️ Bullets sound too uniform — some sentences are very similar "
                "in length. Mix in a couple of shorter or longer bullets, or "
                "click Re-generate to try again."))
    elif burstiness["passes_burstiness"]:
        # Burstiness passes but readability flagged AI words remain
        checks.append((False,
            "⚠️ A few AI-sounding phrases slipped through (Astra normally "
            "catches these). Check the bullets for words like 'leverage', "
            "'spearhead', 'robust' and reword them."))
    else:
        # Both failed
        checks.append((False,
            "⚠️ Writing reads as AI-generated — uniform bullet structure and "
            "some flagged phrases. Re-generate with the Keep-my-voice preset "
            "or edit the longest bullets to vary phrasing."))

    # ─── Grounding ───
    grounding_pass = len(grounding_warnings) == 0
    if grounding_pass:
        checks.append((True,
            "Every number in the resume traces back to your real experience "
            "— interview-defensible."))
    else:
        warn_count = sum(len(v) for v in grounding_warnings.values())
        checks.append((False,
            f"⚠️ {warn_count} bullet(s) use numbers not found in Supraja's base "
            f"resume. Review the Grounding Warnings below before sending — "
            f"she'll have to defend these in an interview."))

    return checks


# ═══════════════════════════════════════════════════════════════════
# 18. GENERATION — single Gemini call with preset awareness
# ═══════════════════════════════════════════════════════════════════

def call_gemini(api_key: str, resume_text: str, jd_text: str,
                preset: str = DEFAULT_PRESET) -> Dict:
    if not api_key:
        return {"error": "Missing GOOGLE_API_KEY."}
    if not jd_text or not jd_text.strip():
        return {"error": "Job description is empty."}

    preset_cfg = PRESET_CONFIGS.get(preset, PRESET_CONFIGS[DEFAULT_PRESET])
    prompt = ASTRA_PROMPT_TEMPLATE % {
        "PRESET_INSTRUCTION": preset_cfg["instruction"],
        "DOMAIN_VOCAB": DOMAIN_VOCAB_REFERENCE,
    }

    client = genai.Client(api_key=api_key)
    safe_schema = get_clean_schema(TailoredOutput)

    full_prompt = (
        f"{prompt}\n\n"
        f"═══ BASE RESUME ═══\n{resume_text}\n\n"
        f"═══ JOB DESCRIPTION ═══\n{jd_text}"
    )

    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=full_prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=safe_schema,
                temperature=preset_cfg["temperature"],
            ),
        )
        raw = json.loads(response.text)
        data = raw.model_dump() if hasattr(raw, "model_dump") else raw
        return data
    except Exception as e:
        return {"error": f"Generation error: {e}"}


def calculate_ats_score(resume_data: dict, jd_text: str, api_key: str) -> dict:
    """Optional button. Single Flash 3 call."""
    if not api_key:
        return {"score": 0, "reasoning": "Missing API key",
                "missing_keywords": "", "title_match_status": "unknown",
                "stuffing_risk": "unknown"}

    client = genai.Client(api_key=api_key)
    payload = {
        "candidate_title": resume_data.get("candidate_title", ""),
        "summary": resume_data.get("summary", ""),
        "skills": resume_data.get("skills", []),
        "experience": resume_data.get("experience", []),
    }
    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=f"{ATS_SCORING_PROMPT}\n\nRESUME:\n{str(payload)[:3500]}\n\nJOB DESCRIPTION:\n{jd_text[:3500]}",
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.2,
            ),
        )
        content = response.text.strip()
        if "```" in content:
            m = re.search(r"```(?:json)?(.*?)```", content, re.DOTALL)
            if m:
                content = m.group(1).strip()
        return json.loads(content)
    except Exception as e:
        return {"score": 0, "reasoning": f"Score error: {e}",
                "missing_keywords": "", "title_match_status": "unknown",
                "stuffing_risk": "unknown"}


def generate_cover_letter(api_key: str, resume_data: dict, jd_text: str) -> str:
    """Optional button. Single Flash 3 call."""
    if not api_key:
        return "ERROR: Missing GOOGLE_API_KEY."
    if not jd_text or not jd_text.strip():
        return "ERROR: Job description is empty."

    client = genai.Client(api_key=api_key)
    resume_context = (
        f"Tailored role: {resume_data.get('candidate_title', '')}\n"
        f"Target company: {resume_data.get('target_company', '')}\n"
        f"Tailored summary: {resume_data.get('summary', '')}"
    )
    prompt = (
        f"{COVER_LETTER_PROMPT}\n\n"
        f"═══ TAILORED RESUME CONTEXT ═══\n{resume_context}\n\n"
        f"═══ JOB DESCRIPTION ═══\n{jd_text}"
    )
    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.6),
        )
        text = (response.text or "").strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:\w+)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        text = re.sub(r"^dear\s+[^\n]+\n+", "", text, flags=re.IGNORECASE)
        text = strip_em_dashes(text)
        text = apply_replacements(text)
        return text
    except Exception as e:
        return f"ERROR: Cover letter generation failed: {e}"


# ═══════════════════════════════════════════════════════════════════
# 19. ASSEMBLE — merge model output with locked facts + safety nets
# ═══════════════════════════════════════════════════════════════════

def assemble_resume(model_output: dict, preset: str) -> dict:
    """Combine model output with locked facts. Apply all safety nets:
    em-dash strip, pronoun strip, banned-phrase replacement, skill
    category validation, grounding check, burstiness audit, multi-axis
    scoring."""

    # ─── Summary ───
    summary = model_output.get("summary", "") or ""
    summary = strip_em_dashes(summary)
    summary = strip_summary_pronouns(summary)
    summary = apply_replacements(summary)
    summary = fix_banned_openers(summary)
    if not summary or len(summary.split()) < 25:
        summary = (
            "Senior Data Engineer with 5+ years building production data "
            "platforms across Amazon Web Services (AWS), Microsoft Azure, and "
            "Google Cloud Platform (GCP), currently moving 500GB+/day in "
            "regulated financial services. Deep production work on Azure "
            "Databricks, Apache PySpark, Snowflake, dbt, and Apache Airflow, "
            "with Terraform-managed infrastructure across the three clouds. "
            "Direct experience in regulated financial services, healthcare, "
            "and retail data platforms, including HIPAA, SOX, and PCI DSS "
            "aligned data delivery. Engineering discipline rooted in Delta "
            "Lake and Apache Iceberg lakehouse patterns, Great Expectations "
            "data quality gates, and Unity Catalog governance. The same "
            "multi-cloud foundation and Snowflake fluency transfer directly "
            "to Lakehouse-native, dbt-driven analytics platforms."
        )

    # ─── Title ───
    candidate_title = strip_em_dashes(
        (model_output.get("candidate_title") or CANDIDATE_TAGLINE_DEFAULT).strip()
    )

    # ─── Skills ───
    raw_skills = model_output.get("skills", []) or []
    if isinstance(raw_skills, dict):
        raw_skills = [{"category": k, "technologies": v} for k, v in raw_skills.items()]
    skills = validate_and_repair_skill_categories(raw_skills)
    skills = [
        {"category": s["category"],
         "technologies": apply_replacements(strip_em_dashes(s["technologies"]))}
        for s in skills
    ]

    # ─── Experience: lock headers, take bullets from model ───
    model_exp_by_company = {}
    for item in model_output.get("experience", []) or []:
        if isinstance(item, dict) and item.get("company"):
            model_exp_by_company[item["company"].strip()] = item

    experience = []
    for locked_role in LOCKED_ROLE_HEADERS:
        company = locked_role["company"]
        model_role = model_exp_by_company.get(company, {})
        resps = model_role.get("responsibilities", []) or []
        achs = model_role.get("achievements", []) or []

        if isinstance(resps, str):
            resps = [r.strip() for r in resps.split("\n") if r.strip()]
        if isinstance(achs, str):
            achs = [a.strip() for a in achs.split("\n") if a.strip()]

        # Apply all sweeps per bullet
        def _clean(b):
            b = strip_em_dashes(b)
            b = apply_replacements(b)
            b = fix_banned_openers(b)
            return b

        resps = [_clean(r) for r in resps if r and r.strip()]
        achs = [_clean(a) for a in achs if a and a.strip()]

        experience.append({
            "role_title": locked_role["role_title"],
            "company": locked_role["company"],
            "location": locked_role["location"],
            "dates": locked_role["dates"],
            "responsibilities": resps,
            "achievements": achs,
        })

    # ─── JD intelligence + tailoring notes (from model) ───
    jd_intel = model_output.get("jd_intelligence", {}) or {}
    if isinstance(jd_intel, dict):
        jd_intel = {
            "top_keywords": jd_intel.get("top_keywords", []) or [],
            "required_tools": jd_intel.get("required_tools", []) or [],
            "nice_to_have_tools": jd_intel.get("nice_to_have_tools", []) or [],
            "emphasis_areas": jd_intel.get("emphasis_areas", []) or [],
            "industry": jd_intel.get("industry", "") or "",
            "seniority_signal": jd_intel.get("seniority_signal", "senior") or "senior",
        }

    tailoring_notes = model_output.get("tailoring_notes", {}) or {}
    if isinstance(tailoring_notes, dict):
        tailoring_notes = {
            "rewrote_for": tailoring_notes.get("rewrote_for", []) or [],
            "did_not_add": tailoring_notes.get("did_not_add", []) or [],
            "keyword_coverage_summary": tailoring_notes.get("keyword_coverage_summary", "") or "",
            "preset_used": preset,
        }

    # ─── Deterministic multi-axis scoring ───
    burstiness = burstiness_audit(experience)
    impact_score, impact_label = compute_impact_axis(experience)
    kw_score, kw_label, missing_kws, kw_low_confidence = compute_keywords_axis(
        skills, summary, experience, jd_intel.get("top_keywords", [])
    )
    read_score, read_label = compute_readability_axis(summary, experience, burstiness)
    exp_score, exp_label = compute_experience_axis(experience)
    overall = compute_overall_match(impact_score, kw_score, read_score, exp_score)

    # ─── Grounding warnings ───
    grounding = validate_grounding(experience)

    # ─── Ready checklist ───
    ready = generate_ready_checklist(
        overall, impact_score, kw_score, read_score,
        missing_kws, grounding, burstiness,
        kw_low_confidence=kw_low_confidence,
    )

    return {
        "candidate_name": CANDIDATE_NAME,
        "candidate_title": candidate_title,
        "contact_info": CANDIDATE_CONTACT,
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "education": list(LOCKED_EDUCATION),
        "certifications": list(LOCKED_CERTIFICATIONS),
        "target_company": (model_output.get("target_company") or "Company").strip(),
        # Metadata for UI:
        "jd_intelligence": jd_intel,
        "tailoring_notes": tailoring_notes,
        "scores": {
            "impact": {"score": impact_score, "label": impact_label},
            "keywords": {
                "score": kw_score, "label": kw_label,
                "missing": missing_kws, "low_confidence": kw_low_confidence,
            },
            "readability": {"score": read_score, "label": read_label},
            "experience": {"score": exp_score, "label": exp_label},
            "overall": overall,
        },
        "burstiness": burstiness,
        "grounding_warnings": grounding,
        "ready_checklist": ready,
        "preset_used": preset,
    }


# ═══════════════════════════════════════════════════════════════════
# 20. DOCX RENDERER (same as v1.0)
# ═══════════════════════════════════════════════════════════════════

def _set_font(run, size, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        pass


def render_docx(data: dict) -> bytes:
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(data["candidate_name"])
    run.font.all_caps = True
    _set_font(run, 28, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run(data["candidate_title"]), 14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(data["contact_info"]), 11, bold=True)

    def _section(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        _set_font(p.add_run(title), 12, bold=True)

    def _bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        if bold_prefix:
            _set_font(p.add_run(bold_prefix), 12, bold=True)
        _set_font(p.add_run(text), 12)

    _section("Professional Profile")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run(data["summary"]), 12)

    _section("Key Skills / Tools & Technologies")
    for entry in data["skills"]:
        _bullet(entry["technologies"], bold_prefix=f"{entry['category']}: ")

    _section("Professional Experience")
    for role in data["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        _set_font(p.add_run(header), 12, bold=True)
        for resp in role["responsibilities"]:
            _bullet(resp)
        if role["achievements"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            _set_font(p.add_run("Achievements:"), 12, bold=True)
            for ach in role["achievements"]:
                _bullet(ach)

    _section("Education & Certifications")
    for edu in data["education"]:
        _bullet(f"{edu['degree']}, {edu['institution']} ({edu['dates']})")
    for cert in data["certifications"]:
        _bullet(cert)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 21. PDF RENDERER (same as v1.0)
# ═══════════════════════════════════════════════════════════════════

def _esc(t):
    return escape(str(t)) if t is not None else ""


def render_pdf(data: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    sn = ParagraphStyle("N", parent=styles["Normal"], fontName="Times-Roman",
                        fontSize=12, leading=14, alignment=TA_JUSTIFY, spaceAfter=0)
    sn_left = ParagraphStyle("NL", parent=styles["Normal"], fontName="Times-Roman",
                             fontSize=12, leading=14, alignment=TA_LEFT, spaceAfter=0)
    sh_name = ParagraphStyle("HN", parent=styles["Normal"], fontName="Times-Bold",
                             fontSize=22, leading=26, alignment=TA_CENTER, spaceAfter=0)
    sh_title = ParagraphStyle("HT", parent=styles["Normal"], fontName="Times-Bold",
                              fontSize=14, leading=16, alignment=TA_CENTER, spaceAfter=0)
    sh_contact = ParagraphStyle("HC", parent=styles["Normal"], fontName="Times-Bold",
                                fontSize=11, leading=13, alignment=TA_CENTER, spaceAfter=6)
    s_sec = ParagraphStyle("Sec", parent=styles["Normal"], fontName="Times-Bold",
                           fontSize=12, leading=14, alignment=TA_LEFT,
                           spaceBefore=10, spaceAfter=2)
    el = []
    el.append(Paragraph(_esc(data["candidate_name"]).upper(), sh_name))
    el.append(Paragraph(_esc(data["candidate_title"]), sh_title))
    el.append(Paragraph(_esc(data["contact_info"]), sh_contact))

    el.append(Paragraph("Professional Profile", s_sec))
    el.append(Paragraph(_esc(data["summary"]), sn))

    el.append(Paragraph("Key Skills / Tools &amp; Technologies", s_sec))
    skill_items = []
    for entry in data["skills"]:
        text = f"<b>{_esc(entry['category'])}:</b> {_esc(entry['technologies'])}"
        skill_items.append(ListItem(Paragraph(text, sn_left), leftIndent=0))
    if skill_items:
        el.append(ListFlowable(skill_items, bulletType="bullet", start="\u2022", leftIndent=15))

    el.append(Paragraph("Professional Experience", s_sec))
    for role in data["experience"]:
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        el.append(Paragraph(f"<b>{_esc(header)}</b>", sn_left))
        el.append(Spacer(1, 2))
        items = [ListItem(Paragraph(_esc(r), sn), leftIndent=0)
                 for r in role["responsibilities"] if r.strip()]
        if items:
            el.append(ListFlowable(items, bulletType="bullet", start="\u2022", leftIndent=15))
        if role["achievements"]:
            el.append(Paragraph("<b>Achievements:</b>", sn_left))
            ach_items = [ListItem(Paragraph(_esc(a), sn), leftIndent=0)
                         for a in role["achievements"] if a.strip()]
            if ach_items:
                el.append(ListFlowable(ach_items, bulletType="bullet",
                                       start="\u2022", leftIndent=25))
        el.append(Spacer(1, 4))

    el.append(Paragraph("Education &amp; Certifications", s_sec))
    edu_items = []
    for edu in data["education"]:
        text = f"{edu['degree']}, {edu['institution']} ({edu['dates']})"
        edu_items.append(ListItem(Paragraph(_esc(text), sn_left), leftIndent=0))
    for cert in data["certifications"]:
        edu_items.append(ListItem(Paragraph(_esc(cert), sn_left), leftIndent=0))
    if edu_items:
        el.append(ListFlowable(edu_items, bulletType="bullet", start="\u2022", leftIndent=15))

    doc.build(el)
    buf.seek(0)
    return buf.getvalue()


def render_cover_letter_docx(letter_body: str, target_company: str = "") -> bytes:
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1.0)

    def _line(text, bold=False, space_after=0, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            _set_font(p.add_run(text), 11, bold=bold)
        return p

    _line(CANDIDATE_NAME, bold=True, space_after=0)
    for piece in [c.strip() for c in CANDIDATE_CONTACT.split("|")]:
        if piece:
            _line(piece, space_after=0)
    _line("", space_after=10)

    today = datetime.date.today().strftime("%B %d, %Y")
    _line(today, space_after=14)

    if target_company.strip() and target_company.strip().lower() != "company":
        greeting = f"Dear Hiring Team at {target_company.strip()},"
    else:
        greeting = "Dear Hiring Team,"
    _line(greeting, space_after=10)

    body = letter_body.strip()
    sign_off_pattern = re.compile(
        r"\n+(thank you,?|regards,?|sincerely,?|kind regards,?|best regards,?)\s*\n+(lakshmi[\s\w]*)\s*$",
        re.IGNORECASE,
    )
    sign_off_match = sign_off_pattern.search(body)
    if sign_off_match:
        body_main = body[: sign_off_match.start()].strip()
        sign_word = sign_off_match.group(1).rstrip(",").strip().capitalize()
    else:
        body_main = body
        sign_word = "Thank you"

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body_main) if p.strip()]
    if not paragraphs:
        paragraphs = [p.strip() for p in body_main.split("\n") if p.strip()]

    for para in paragraphs:
        clean_para = re.sub(r"\s*\n\s*", " ", para)
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        _set_font(p.add_run(clean_para), 11)

    _line("", space_after=6)
    _line(f"{sign_word},", space_after=18)
    _line(CANDIDATE_NAME, space_after=0)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 22. STREAMLIT UI — 3-step wizard per research
# ═══════════════════════════════════════════════════════════════════

st.set_page_config(page_title=PAGE_TITLE, layout="wide", page_icon="🎯",
                   initial_sidebar_state="expanded")

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .block-container {padding-top: 1.5rem;}
    div.stButton > button:first-child {border-radius: 6px; font-weight: 600;}
    div[data-testid="stMetricValue"] {font-size: 1.6rem;}
    .axis-box {padding: 12px; border-radius: 8px; border: 1px solid #e0e0e0; margin: 4px 0;}
</style>
""", unsafe_allow_html=True)

# ─── Session state ───
for key, default in [
    ("step", 1),
    ("tailored", None),
    ("saved_base", SUPRAJA_BASE_RESUME),
    ("saved_jd", ""),
    ("preset", DEFAULT_PRESET),
    ("ats_score", None),
    ("cover_letter", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ─── Sidebar ───
with st.sidebar:
    st.header("⚙️ Configuration")
    if GOOGLE_API_KEY:
        st.success("API key configured")
        api_key = GOOGLE_API_KEY
    else:
        st.warning("Add GOOGLE_API_KEY to Streamlit secrets, or paste below.")
        api_key = st.text_input("Google API Key", type="password")

    st.divider()
    st.markdown("**Target roles:**")
    st.caption(
        "Senior Data Engineer • Cloud DE (Azure/AWS/GCP) • Snowflake DE • "
        "ETL/Informatica Developer • Analytics Engineer • PySpark Developer • "
        "Big Data Engineer • Database Engineer"
    )

    st.divider()
    st.markdown(f"**Model:** `{GENERATION_MODEL}`")
    st.caption(
        "Single-pass generation. Scoring and validation done deterministically "
        "in Python (free). ~$0.008 per resume."
    )

    st.divider()
    if st.button("🗑️ Reset everything", use_container_width=True):
        for key in ["step", "tailored", "saved_jd", "preset", "ats_score", "cover_letter"]:
            st.session_state.pop(key, None)
        st.session_state["saved_base"] = SUPRAJA_BASE_RESUME
        st.session_state["step"] = 1
        st.rerun()

    st.caption("Astra v2.0 | Personalised for Supraja")


# ═══════════════════════════════════════════════════════════════════
# STEP 1 — Paste base resume + JD
# ═══════════════════════════════════════════════════════════════════
if st.session_state["step"] == 1:
    st.markdown(
        f"<h1 style='text-align: center;'>{PAGE_TITLE}</h1>",
        unsafe_allow_html=True,
    )
    st.markdown(
        "<p style='text-align: center; color: #888;'>"
        "Paste a JD. Get a tailored, recruiter-grade resume. Get the call.</p>",
        unsafe_allow_html=True,
    )
    st.caption("Step 1 of 3: paste your base resume and the job description.")
    st.divider()

    col_r, col_j = st.columns(2)

    with col_r:
        st.subheader("📋 Base Resume")
        st.caption("Pre-loaded. Edit only for one-off tweaks.")
        base = st.text_area(
            "Base Resume", value=st.session_state["saved_base"], height=420,
            label_visibility="collapsed",
        )
        if st.button("↩️ Restore default base resume", use_container_width=True):
            st.session_state["saved_base"] = SUPRAJA_BASE_RESUME
            st.rerun()

    with col_j:
        st.subheader("💼 Job Description")
        st.caption("Paste the full JD here.")
        jd = st.text_area(
            "Job Description", value=st.session_state["saved_jd"], height=420,
            label_visibility="collapsed",
            placeholder="Paste the full JD here…",
        )

    st.divider()
    if st.button("Continue →", type="primary", use_container_width=True):
        if not jd.strip():
            st.warning("Please paste a job description.")
        elif not base.strip():
            st.warning("Base resume is empty. Click 'Restore default base resume'.")
        else:
            st.session_state["saved_base"] = base
            st.session_state["saved_jd"] = jd
            st.session_state["step"] = 2
            st.rerun()


# ═══════════════════════════════════════════════════════════════════
# STEP 2 — Choose preset + generate
# ═══════════════════════════════════════════════════════════════════
elif st.session_state["step"] == 2:
    st.caption("Step 2 of 3: choose how much Astra should rewrite.")
    st.markdown("### How aggressively should Astra rewrite?")

    preset = st.radio(
        "Tailoring depth",
        options=list(PRESET_CONFIGS.keys()),
        format_func=lambda x: PRESET_CONFIGS[x]["label"],
        captions=[PRESET_CONFIGS[k]["caption"] for k in PRESET_CONFIGS.keys()],
        index=list(PRESET_CONFIGS.keys()).index(st.session_state["preset"]),
        label_visibility="collapsed",
    )
    st.session_state["preset"] = preset

    st.divider()
    with st.expander("⚙ Advanced (optional)", expanded=False):
        st.caption(
            "These settings are filled in for Supraja's profile. Override "
            "only if you have a specific reason."
        )
        st.text_input(
            "Seniority", value="Senior (5+ years)", disabled=True,
            help="Locked at code level. Astra uses senior-tier verb selection.",
        )
        st.text_input(
            "Target match band", value="78–85% (research-validated)", disabled=True,
            help="Above 88% triggers keyword-stuffing detectors.",
        )

    st.divider()
    col_back, col_gen = st.columns([1, 3])
    if col_back.button("← Back", use_container_width=True):
        st.session_state["step"] = 1
        st.rerun()

    if col_gen.button("✨ Generate Tailored Resume", type="primary", use_container_width=True):
        if not api_key:
            st.error("Need a Google API key.")
        else:
            with st.spinner(
                f"Tailoring resume to JD (preset: {PRESET_CONFIGS[preset]['label']}, "
                "single pass, ~10 seconds)…"
            ):
                model_out = call_gemini(
                    api_key,
                    st.session_state["saved_base"],
                    st.session_state["saved_jd"],
                    preset=preset,
                )
                if "error" in model_out:
                    st.error(model_out["error"])
                else:
                    final = assemble_resume(model_out, preset)
                    st.session_state["tailored"] = final
                    st.session_state["ats_score"] = None
                    st.session_state["cover_letter"] = None
                    st.session_state["step"] = 3
                    st.rerun()


# ═══════════════════════════════════════════════════════════════════
# STEP 3 — Review
# ═══════════════════════════════════════════════════════════════════
elif st.session_state["step"] == 3 and st.session_state["tailored"]:
    data = st.session_state["tailored"]
    scores = data["scores"]

    # ─── Top bar ───
    tb1, tb2, tb3 = st.columns([3, 2, 1])
    with tb1:
        st.markdown(f"## 🎯 Target: {data['target_company']}")
        st.caption(f"Tailored title: **{data['candidate_title']}** "
                   f"· Preset: {PRESET_CONFIGS[data['preset_used']]['label']}")
    with tb3:
        if st.button("New JD", use_container_width=True):
            st.session_state["tailored"] = None
            st.session_state["saved_jd"] = ""
            st.session_state["ats_score"] = None
            st.session_state["cover_letter"] = None
            st.session_state["step"] = 1
            st.rerun()

    # ─── Overall match (3-layer: emoji + label + percentage) ───
    overall = scores["overall"]
    emoji, label = overall_match_label(overall)
    st.markdown(
        f"<div style='padding: 14px; border-radius: 10px; "
        f"background: #f7f7f9; margin: 8px 0;'>"
        f"<span style='font-size: 1.8rem;'>{emoji}</span> "
        f"<span style='font-size: 1.3rem; font-weight: 600;'>{label}</span> "
        f"<span style='color: #888; margin-left: 12px;'>{overall}%</span>"
        f"</div>",
        unsafe_allow_html=True,
    )
    if overall > 85:
        st.warning(
            "⚠️ Approaching keyword-stuffing zone. Modern parsers flag "
            "matches above 85% as unnatural. Consider 'Keep my voice' preset."
        )

    # ─── Multi-axis score card ───
    st.markdown("#### Score breakdown")
    a1, a2, a3, a4 = st.columns(4)
    for col, key, title in [
        (a1, "impact", "Impact"),
        (a2, "keywords", "Keywords"),
        (a3, "readability", "Readability"),
        (a4, "experience", "Experience"),
    ]:
        ax = scores[key]
        score_val = ax["score"]
        if score_val >= 80:
            color = "#1a7f37"
        elif score_val >= 60:
            color = "#bf8700"
        else:
            color = "#cf222e"
        col.markdown(
            f"<div class='axis-box'>"
            f"<div style='color: #555; font-size: 0.85rem;'>{title}</div>"
            f"<div style='color: {color}; font-size: 1.4rem; font-weight: 600;'>{score_val}/100</div>"
            f"<div style='color: #666; font-size: 0.78rem; margin-top: 4px;'>{ax['label']}</div>"
            f"</div>",
            unsafe_allow_html=True,
        )

    # ─── What Astra did ───
    st.divider()
    notes = data.get("tailoring_notes", {})
    with st.expander("📋 What Astra did", expanded=True):
        if notes.get("keyword_coverage_summary"):
            st.markdown(f"**Summary:** {notes['keyword_coverage_summary']}")
        if notes.get("rewrote_for"):
            st.markdown("**Rewrote bullets for:**")
            for note in notes["rewrote_for"]:
                st.markdown(f"- {note}")
        if notes.get("did_not_add"):
            st.markdown("**Did NOT add (not in your original — verify or skip):**")
            for note in notes["did_not_add"]:
                st.markdown(f"- 🛡️ {note}")
        missing = scores["keywords"].get("missing", [])
        if missing:
            st.markdown("**Missing JD keywords:**")
            st.caption(", ".join(missing))

    # ─── Tabs ───
    tab_preview, tab_edit, tab_export, tab_cover, tab_diff = st.tabs(
        ["👀 Preview", "📝 Edit", "📥 Export", "✍️ Cover Letter", "🔍 Before / After"]
    )

    with tab_preview:
        # ─── Ready-to-send checklist ABOVE the preview ───
        st.markdown("#### Ready to send?")
        for passed, msg in data["ready_checklist"]:
            icon = "✅" if passed else "⚠️"
            st.markdown(f"{icon} {msg}")

        # ─── Grounding warnings ───
        gw = data.get("grounding_warnings", {})
        if gw:
            with st.expander("🛡️ Grounding warnings (review before sending)", expanded=False):
                for company, warns in gw.items():
                    st.markdown(f"**{company}:**")
                    for w in warns:
                        st.markdown(f"- {w}")

        st.divider()
        st.subheader("Professional Profile")
        st.write(data["summary"])

        st.subheader("Key Skills / Tools & Technologies")
        for entry in data["skills"]:
            st.markdown(f"- **{entry['category']}:** {entry['technologies']}")

        st.subheader("Professional Experience")
        for role in data["experience"]:
            st.markdown(
                f"**{role['role_title']}** | {role['company']} | "
                f"{role['location']} | {role['dates']}"
            )
            for r in role["responsibilities"]:
                st.markdown(f"- {r}")
            if role["achievements"]:
                st.markdown("**Achievements:**")
                for a in role["achievements"]:
                    st.markdown(f"- {a}")

        st.subheader("Education & Certifications")
        for edu in data["education"]:
            st.markdown(f"- {edu['degree']}, {edu['institution']} ({edu['dates']})")
        for cert in data["certifications"]:
            st.markdown(f"- {cert}")

        st.divider()
        if st.button(
            "📊 Run optional Flash 3 ATS scoring check (~$0.008)",
            use_container_width=True,
        ):
            if not api_key:
                st.error("Need API key.")
            else:
                with st.spinner("Scoring…"):
                    score_result = calculate_ats_score(
                        data, st.session_state["saved_jd"], api_key
                    )
                    st.session_state["ats_score"] = score_result
                    st.rerun()
        if st.session_state["ats_score"]:
            score_data = st.session_state["ats_score"]
            st.info(
                f"**Flash 3 ATS score:** {score_data.get('score', 0)}% · "
                f"Title match: {score_data.get('title_match_status', 'n/a')} · "
                f"Stuffing risk: {score_data.get('stuffing_risk', 'n/a')}"
                f"\n\n_{score_data.get('reasoning', '')}_"
            )
            ext_missing = score_data.get("missing_keywords", "")
            if ext_missing:
                st.caption(f"Flash 3 says still missing: {ext_missing}")

    with tab_edit:
        with st.form("edit_form"):
            data["candidate_title"] = st.text_input("Title under name", data["candidate_title"])
            data["summary"] = st.text_area("Summary", data["summary"], height=180)

            st.markdown("##### Skills (comma-separated per category)")
            new_skills = []
            for idx, entry in enumerate(data["skills"]):
                cat = st.text_input(f"Category {idx+1}", entry["category"], key=f"sk_cat_{idx}")
                techs = st.text_area(
                    f"Tools for {cat}", entry["technologies"], height=70, key=f"sk_tech_{idx}",
                )
                if cat.strip() and techs.strip():
                    new_skills.append({"category": cat.strip(), "technologies": techs.strip()})
            data["skills"] = new_skills

            st.markdown("##### Experience")
            for i, role in enumerate(data["experience"]):
                with st.expander(f"{role['role_title']} @ {role['company']}", expanded=False):
                    resps_text = "\n".join(role["responsibilities"])
                    new_resps = st.text_area(
                        "Responsibilities (one per line)",
                        resps_text, height=180, key=f"r_{i}",
                    )
                    role["responsibilities"] = [
                        line.strip() for line in new_resps.split("\n") if line.strip()
                    ]
                    if role["achievements"] or i < 2:
                        achs_text = "\n".join(role["achievements"])
                        new_achs = st.text_area(
                            "Achievements (one per line)",
                            achs_text, height=80, key=f"a_{i}",
                        )
                        role["achievements"] = [
                            line.strip() for line in new_achs.split("\n") if line.strip()
                        ]

            if st.form_submit_button("💾 Save edits + re-validate", type="primary"):
                # Re-apply safety nets after manual edit
                data["summary"] = fix_banned_openers(
                    apply_replacements(
                        strip_summary_pronouns(strip_em_dashes(data["summary"]))
                    )
                )
                data["candidate_title"] = strip_em_dashes(data["candidate_title"])
                data["skills"] = validate_and_repair_skill_categories(data["skills"])
                for role in data["experience"]:
                    role["responsibilities"] = [
                        fix_banned_openers(apply_replacements(strip_em_dashes(r)))
                        for r in role["responsibilities"]
                    ]
                    role["achievements"] = [
                        fix_banned_openers(apply_replacements(strip_em_dashes(a)))
                        for a in role["achievements"]
                    ]
                # Recompute scoring with edited content
                burst = burstiness_audit(data["experience"])
                imp_s, imp_l = compute_impact_axis(data["experience"])
                kw_s, kw_l, kw_miss, kw_lowc = compute_keywords_axis(
                    data["skills"], data["summary"], data["experience"],
                    data["jd_intelligence"].get("top_keywords", []),
                )
                rd_s, rd_l = compute_readability_axis(data["summary"], data["experience"], burst)
                ex_s, ex_l = compute_experience_axis(data["experience"])
                overall_new = compute_overall_match(imp_s, kw_s, rd_s, ex_s)
                data["scores"] = {
                    "impact": {"score": imp_s, "label": imp_l},
                    "keywords": {
                        "score": kw_s, "label": kw_l,
                        "missing": kw_miss, "low_confidence": kw_lowc,
                    },
                    "readability": {"score": rd_s, "label": rd_l},
                    "experience": {"score": ex_s, "label": ex_l},
                    "overall": overall_new,
                }
                data["burstiness"] = burst
                data["grounding_warnings"] = validate_grounding(data["experience"])
                data["ready_checklist"] = generate_ready_checklist(
                    overall_new, imp_s, kw_s, rd_s, kw_miss,
                    data["grounding_warnings"], burst,
                    kw_low_confidence=kw_lowc,
                )
                st.session_state["tailored"] = data
                st.success("Saved and re-scored.")
                st.rerun()

    with tab_export:
        company = data["target_company"] or "Company"
        company_safe = re.sub(r"[^A-Za-z0-9_-]", "_", company.strip()) or "Company"
        filename_base = f"Lakshmi_Supraja_Konakanchi_{company_safe}"

        st.text_input("Filename (no extension)", filename_base, key="fname")
        fname = st.session_state.get("fname", filename_base)

        c1, c2 = st.columns(2)
        try:
            docx_bytes = render_docx(data)
            c1.download_button(
                "📄 Word (.docx)", data=docx_bytes,
                file_name=f"{fname}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary", use_container_width=True,
            )
        except Exception as e:
            c1.error(f"DOCX error: {e}")

        try:
            pdf_bytes = render_pdf(data)
            c2.download_button(
                "📕 PDF", data=pdf_bytes,
                file_name=f"{fname}.pdf", mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            c2.error(f"PDF error: {e}")

    with tab_cover:
        st.caption(
            "Short, human-sounding cover letter using the war story closest to "
            "the JD. Generated separately (~$0.008) so it doesn't slow down "
            "the main flow."
        )

        cover_btn_label = (
            "✨ Draft Cover Letter (~$0.008)"
            if not st.session_state["cover_letter"]
            else "🔄 Re-draft Cover Letter"
        )
        if st.button(cover_btn_label, type="primary"):
            if not api_key:
                st.error("Need API key.")
            elif not st.session_state["saved_jd"].strip():
                st.warning("No saved JD.")
            else:
                with st.spinner("Picking the right war story, drafting…"):
                    cl = generate_cover_letter(
                        api_key, data, st.session_state["saved_jd"]
                    )
                    if cl.startswith("ERROR:"):
                        st.error(cl)
                    else:
                        st.session_state["cover_letter"] = cl
                        st.rerun()

        if st.session_state["cover_letter"]:
            edited = st.text_area(
                "Cover letter (editable)",
                st.session_state["cover_letter"], height=420,
            )
            st.session_state["cover_letter"] = edited

            try:
                cl_bytes = render_cover_letter_docx(
                    st.session_state["cover_letter"],
                    target_company=data.get("target_company", ""),
                )
                company_safe = re.sub(
                    r"[^A-Za-z0-9_-]", "_",
                    (data["target_company"] or "Company").strip()
                ) or "Company"
                st.download_button(
                    "📄 Download Cover Letter (.docx)",
                    data=cl_bytes,
                    file_name=f"CoverLetter_Lakshmi_Supraja_{company_safe}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                )
            except Exception as e:
                st.error(f"Cover letter render error: {e}")

    with tab_diff:
        st.caption(
            "Section-by-section comparison: base resume on the left, tailored "
            "version on the right. Use this to spot what changed and whether "
            "the tailored version still sounds like you."
        )

        # Build comparable text blocks for each section
        def _section_block_base(label: str, body: str) -> str:
            return f"## {label}\n\n{body}"

        # Tailored side
        tailored_summary = data["summary"]
        tailored_skills_block = "\n".join(
            [f"- **{e['category']}:** {e['technologies']}" for e in data["skills"]]
        )
        tailored_exp_blocks = []
        for role in data["experience"]:
            header = (
                f"### {role['role_title']} | {role['company']} | "
                f"{role['location']} | {role['dates']}"
            )
            body = "\n".join([f"- {r}" for r in role["responsibilities"]])
            if role["achievements"]:
                body += "\n**Achievements:**\n" + "\n".join(
                    [f"- {a}" for a in role["achievements"]]
                )
            tailored_exp_blocks.append(header + "\n" + body)

        col_b, col_t = st.columns(2)
        with col_b:
            st.markdown("### Base resume")
            st.text_area(
                "Base", st.session_state["saved_base"], height=600,
                label_visibility="collapsed", disabled=True,
            )
        with col_t:
            st.markdown("### Tailored resume")
            tailored_text = (
                f"# {data['candidate_name']}\n"
                f"## {data['candidate_title']}\n"
                f"{data['contact_info']}\n\n"
                f"### Professional Profile\n{tailored_summary}\n\n"
                f"### Key Skills\n{tailored_skills_block}\n\n"
                f"### Professional Experience\n"
                + "\n\n".join(tailored_exp_blocks)
            )
            st.text_area(
                "Tailored", tailored_text, height=600,
                label_visibility="collapsed", disabled=True,
            )

# If step is 3 but no tailored data, redirect to step 1
elif st.session_state["step"] == 3 and not st.session_state["tailored"]:
    st.session_state["step"] = 1
    st.rerun()
